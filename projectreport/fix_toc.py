"""Add missing TOC entries for new sections."""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

SRC = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx'
doc = Document(SRC)

def insert_after(para, text):
    """Insert a body text paragraph after para."""
    ref = para._element if hasattr(para, '_element') else para
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    new_p.append(r)
    # Copy style
    pPr = ref.find(qn('w:pPr'))
    if pPr is not None:
        new_pPr = deepcopy(pPr)
        pStyle = new_pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            sid = doc.styles['Body Text'].element.get(qn('w:styleId'))
            pStyle.set(qn('w:val'), sid)
        new_p.insert(0, new_pPr)
    ref.addnext(new_p)
    return new_p

def find_para(text_fragment):
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text and p.text.strip():
            return i, p
    return None, None

# 1. Add Data Dictionary, ER Diagram, DFD after "Use Case Diagram" (P34)
_, p_uc = find_para('Use Case Diagram')
insert_after(p_uc, '\t\t\t\t\t\t\t\tData Dictionary')
insert_after(p_uc, '\t\t\t\t\t\t\tEntity Relationship Diagram')
insert_after(p_uc, '\t\t\t\t\t\t\t\tData Flow Diagram')
print("Added Data Dictionary/ER/DFD to TOC")

# 2. Add "Implementation of Security" after "Input Output Screens"
_, p_io = find_para('Input Output Screens')
insert_after(p_io, 'Implementation of Security for the Software Developed')
print("Added Security chapter to TOC")

# 3. Fix numbering in TOC
renames_toc = {
    'Implementation Plan': '\t\t\t\t\t\t\t9. Implementation Plan',
    'Limitations of the Project': '\t\t\t\t\t10. Limitations of the Project',
    'Future Application of the Project': '\t\t\t11. Future Application of the Project',
    'Conclusion': '\t\t\t\t\t\t\t12. Conclusion',
    'Bibliography': '\t\t\t\t\t\t\t13. Bibliography',
}

for p in doc.paragraphs:
    t = p.text.strip()
    # Remove old numbers (like "11. Future Application...") and replace
    for key, new_val in renames_toc.items():
        if key in t and not t.startswith(key.replace('9.', '').replace('10.', '').replace('11.', '').replace('12.', '').replace('13.', '')):
            # Check if it's a numbered entry
            parts = t.split('\t')
            for part in parts:
                if key in part:
                    # Found it - replace the whole text
                    for run in p.runs:
                        run.text = ''
                    if p.runs:
                        p.runs[0].text = new_val
                    else:
                        r = OxmlElement('w:r')
                        t_el = OxmlElement('w:t')
                        t_el.text = new_val
                        t_el.set(qn('xml:space'), 'preserve')
                        r.append(t_el)
                        p._element.append(r)

# More direct approach - find each by exact text and update
print("TOC entries updated!")

doc.save(SRC)
print("Saved!")
