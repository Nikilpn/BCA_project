"""Reorder sections in Chapter 3 of nikhilproject241.docx.

Current physical order (top to bottom):
  3.4 Software and Hardware Requirements
  3.3 Use Case Diagram
  3.7 Data Flow Diagram
  3.6 Entity Relationship Diagram
  3.5 Data Dictionary

Desired physical order:
  3.3 Use Case Diagram
  3.4 Software and Hardware Requirements
  3.5 Data Dictionary
  3.6 Entity Relationship Diagram
  3.7 Data Flow Diagram
"""
import docx
from lxml import etree
from docx.oxml.ns import qn

doc = docx.Document('/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx')

body_elem = doc.element.body

# First, identify the ranges for each section
sections = {}  # heading_text -> (start_idx, end_idx) in doc.paragraphs

# Find all headings in Ch 3
sec_starts = {}
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    s = p.style.name
    if 'Heading' in s:
        if t.startswith('3.3') or t.startswith('3.4') or t.startswith('3.5') or t.startswith('3.6') or t.startswith('3.7'):
            sec_starts[t] = i
        elif t.startswith('4.'):
            break  # stop at Ch 4

print("Section starts found:")
for k, v in sec_starts.items():
    print(f"  P{v}: {k}")

# Build ranges: heading_para -> (start_p_elem, end_p_elem_exclusive)
# The range includes all paragraphs from the heading up to (but not including) the next heading or Ch 4
sorted_starts = sorted(sec_starts.items(), key=lambda x: x[1])
ranges = []
for idx, (heading, start_idx) in enumerate(sorted_starts):
    if idx + 1 < len(sorted_starts):
        end_idx = sorted_starts[idx+1][1]
    else:
        # Last section before Ch 4
        for j in range(start_idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[j].text.strip().startswith('4. Design Document') and 'Heading' in doc.paragraphs[j].style.name:
                end_idx = j
                break
        else:
            end_idx = len(doc.paragraphs)
    
    # Get the paragraph elements
    start_elem = doc.paragraphs[start_idx]._element
    end_elem = doc.paragraphs[end_idx]._element if end_idx < len(doc.paragraphs) else None
    ranges.append((heading, start_elem, end_elem, doc.paragraphs[start_idx]))
    print(f"  Section '{heading}': P{start_idx}-{end_idx-1}")

# Collect elements for each section
def get_section_elements(start_elem, end_elem):
    """Get all paragraph elements from start_elem up to (but not including) end_elem."""
    elems = []
    current = start_elem
    while current is not None and (end_elem is None or current != end_elem):
        if current.tag == qn('w:p'):
            elems.append(current)
        # Check next sibling
        if current.tail is not None and current.tail.strip():
            pass
        current = current.getnext()
        if current is None or (end_elem is not None and current == end_elem):
            break
    return elems

# Get elements for each section
section_elems = []
for heading, start_elem, end_elem, para_obj in ranges:
    elems = get_section_elements(start_elem, end_elem)
    section_elems.append((heading, elems))
    print(f"  Section '{heading}': {len(elems)} paragraph elements")

# Determine the correct order based on section numbers
# Current order: 3.4, 3.3, 3.7, 3.6, 3.5
# Desired order: 3.3, 3.4, 3.5, 3.6, 3.7
section_order = sorted(section_elems, key=lambda x: x[0])

print("\nReordering sections:")
for heading, elems in section_order:
    print(f"  '{heading}' -> {len(elems)} elements")

# Find the insertion point: right after the last element of the first section currently
# The first section in document body is 3.4. Its elements are at the start of Ch 3 content.
# We want to move everything into: 3.3, 3.4, 3.5, 3.6, 3.7

# Remove all section elements from their current positions, then re-insert in order
# Get all section elements as a flat list
all_elems_to_move = []
for _, elems in section_order:
    for e in elems:
        all_elems_to_move.append(e)

# All these elements are already in the document, just in wrong order.
# We need to physically reorder them.

# Strategy: detach all elements, then re-insert at the position of the first one
first_elem = all_elems_to_move[0]
parent = first_elem.getparent()

# Find the previous sibling of the first element (to insert after it)
prev_sibling = first_elem.getprevious()

# Detach all elements from the tree
for e in all_elems_to_move:
    parent.remove(e)

# Now re-insert in the correct order
# Insert after prev_sibling (or as first child if no prev sibling)
insert_point = prev_sibling
for e in all_elems_to_move:
    if insert_point is None:
        # Insert as first child
        parent.insert(0, e)
    else:
        insert_point.addnext(e)
    insert_point = e

print("Reordering complete!")

output_path = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx'
doc.save(output_path)
print(f"Document saved: {output_path}")
