"""Step 1: Fix 3.4/3.3 ordering. Save. Step 2: Add 3.5/3.6/3.7. Step 3: Add Security + renumber."""
import sys, os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

SRC = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx'

def add_after(para, text, style_name=None):
    ref = para._element if hasattr(para, '_element') else para
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    new_p.append(r)
    if style_name:
        pPr = ref.find(qn('w:pPr'))
        if pPr is not None:
            new_pPr = deepcopy(pPr)
            pStyle = new_pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                sid = doc.styles[style_name].element.get(qn('w:styleId'))
                pStyle.set(qn('w:val'), sid)
            new_p.insert(0, new_pPr)
    ref.addnext(new_p)
    return new_p

def body_after(p, t): return add_after(p, t, 'Body Text')
def h2_after(p, t):  return add_after(p, t, 'Heading 2.section')
def h1_after(p, t):  return add_after(p, t, 'Heading 1.chapter')

def find_heading(text_fragment):
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        s = p.style.name
        if text_fragment in t and 'Heading' in s:
            return i, p
    return None, None

def last_body_before(heading_text):
    """Return the last non-empty paragraph before the given heading text."""
    hidx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        s = p.style.name
        if t.startswith(heading_text) and 'Heading' in s:
            hidx = i
            break
    if hidx is None: return None
    for i in range(hidx - 1, -1, -1):
        t = doc.paragraphs[i].text.strip()
        if t:
            return doc.paragraphs[i]
    return None

def print_headings(label):
    print(f"\n=== {label} ===")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        s = p.style.name
        if 'Heading' in s and t and (t[0].isdigit() or t.startswith('3.')):
            print(f"  P{i}: {t[:90]}")

# ============================================================
# PHASE 1: Fix 3.4/3.3 ordering via XML element manipulation
# ============================================================
doc = Document(SRC)
BODY_EL = doc.element.body

_, p_34 = find_heading('3.4 Software and Hardware')
_, p_33 = find_heading('3.3 Use Case')

# Collect 3.4 elements (from 3.4 heading to before 3.3 heading)
el_34 = []
cur = p_34._element
while cur is not None and cur != p_33._element:
    if cur.tag == qn('w:p'):
        el_34.append(cur)
    cur = cur.getnext()

# Remove from current position
for e in el_34:
    BODY_EL.remove(e)

# Insert after last element of 3.3 section  
# Find element to insert after - we need the last para of 3.3
# Get everything from 3.3 heading to before 4. Design heading
_, p_4 = find_heading('4. Design Document')
last_33 = p_33._element
cur = p_33._element.getnext()
while cur is not None and cur != p_4._element:
    if cur.tag == qn('w:p'):
        last_33 = cur
    cur = cur.getnext()

# Insert 3.4 elements after last_33
for e in el_34:
    last_33.addnext(e)
    last_33 = e

print("Phase 1: 3.4 moved after 3.3")
print_headings("After Phase 1")

# Save intermediate
doc.save(SRC)
doc = Document(SRC)  # Reopen

# ============================================================
# PHASE 2: Add 3.5 Data Dictionary, 3.6 ER, 3.7 DFD
# ============================================================
anchor = last_body_before('4. Design Document')
print(f"\nPhase 2 anchor: {anchor.text.strip()[:60]}")

# 3.5 Data Dictionary
el = h2_after(anchor, "3.5 Data Dictionary")
el = body_after(el, "The data dictionary defines the structure, data types, constraints, and descriptions for all database tables.")

el = body_after(el, "Table 3.1 - roomtypedb (Room Category)")
el = body_after(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body_after(el, "ROOMTYPE | VARCHAR(100) | NOT NULL | Category name (Luxury, Deluxe, Standard, Budget)")
el = body_after(el, "ROOMTYPEIMAGE | VARCHAR(255) | NULLABLE | Image file path")
el = body_after(el, "DESCRIPTION | TEXT | NULLABLE | Category description")

el = body_after(el, "Table 3.2 - roomnamedb (Room)")
el = body_after(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body_after(el, "ROOMTYPE | Integer | FK (roomtypedb.id) | Room category reference")
el = body_after(el, "ROOMNAME | VARCHAR(100) | NOT NULL | Room name/number")
el = body_after(el, "ROOMIMAGE | VARCHAR(255) | NULLABLE | Image file path")
el = body_after(el, "ROOMPRICE | Integer | NOT NULL | Price per night (INR)")
el = body_after(el, "ROOMDESCRIPTION | TEXT | NULLABLE | Room features")

el = body_after(el, "Table 3.3 - staffdb (Staff Profile)")
el = body_after(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body_after(el, "STAFFNAME | VARCHAR(100) | NOT NULL | Staff full name")
el = body_after(el, "DESIGNATION | VARCHAR(100) | NOT NULL | Job title")
el = body_after(el, "FACEBOOKURL | VARCHAR(255) | NULLABLE | Facebook URL")
el = body_after(el, "INSTAGRAMURL | VARCHAR(255) | NULLABLE | Instagram URL")
el = body_after(el, "STAFFIMAGE | VARCHAR(255) | NULLABLE | Photograph path")

el = body_after(el, "Table 3.4 - Registerdb (Customer Registration)")
el = body_after(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body_after(el, "USERNAME | VARCHAR(100) | UNIQUE, NOT NULL | Login username")
el = body_after(el, "PASSWORD | VARCHAR(100) | NOT NULL | Password")
el = body_after(el, "CONFIRMPASSWORD | VARCHAR(100) | NOT NULL | Password confirmation")
el = body_after(el, "EMAIL | VARCHAR(100) | NOT NULL | Email address")

el = body_after(el, "Table 3.5 - bookingdb (Booking Record)")
el = body_after(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body_after(el, "CUSTOMER | Integer | FK (Registerdb.id), SET_NULL | Customer reference")
el = body_after(el, "CUSTOMERNAME | VARCHAR(100) | NOT NULL | Customer name")
el = body_after(el, "CONTACTEMAIL | VARCHAR(100) | NOT NULL | Customer email")
el = body_after(el, "CHECKIN | Date | NOT NULL | Check-in date")
el = body_after(el, "CHECKOUT | Date | NOT NULL | Check-out date")
el = body_after(el, "TOTALADULTS | Integer | NOT NULL, DEFAULT 1 | Adult count")
el = body_after(el, "TOTALCHILDS | Integer | NOT NULL, DEFAULT 0 | Child count")
el = body_after(el, "SELECTROOM | Integer | FK (roomnamedb.id), CASCADE | Room reference")
el = body_after(el, "SPECIALREQUEST | TEXT | NULLABLE | Special requests")
el = body_after(el, "TOTALPRICE | Integer | NOT NULL | Total price (INR)")

el = body_after(el, "Table 3.6 - Totaldb (Payment Transaction)")
el = body_after(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body_after(el, "BOOKING | Integer | FK (bookingdb.id), CASCADE | Booking reference")
el = body_after(el, "CUSTOMERNAME | VARCHAR(100) | NOT NULL | Customer name")
el = body_after(el, "MOBILE | VARCHAR(15) | NOT NULL | Mobile number")
el = body_after(el, "TOTALPRICE | Integer | NOT NULL | Amount paid (INR)")

el = body_after(el, "Table 3.7 - customercontactdb (Contact Enquiry)")
el = body_after(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body_after(el, "CONTACTNAME | VARCHAR(100) | NOT NULL | Sender name")
el = body_after(el, "CONTACTNUMBER | VARCHAR(15) | NULLABLE | Phone number")
el = body_after(el, "CONTACTEMAIL | VARCHAR(100) | NOT NULL | Sender email")
el = body_after(el, "CONTACTSUBJECT | VARCHAR(200) | NOT NULL | Subject")
el = body_after(el, "CONTACTMESSAGE | TEXT | NOT NULL | Message content")

el = body_after(el, "Table 3.8 - ChatbotResponse")
el = body_after(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body_after(el, "keywords | TEXT | NOT NULL | Keywords for pattern matching")
el = body_after(el, "response | TEXT | NOT NULL | Bot response text")
el = body_after(el, "query_type | VARCHAR(50) | NOT NULL | Query category")

el = body_after(el, "Table 3.9 - ChatbotConversation")
el = body_after(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body_after(el, "username | VARCHAR(100) | NOT NULL | Participant username")
el = body_after(el, "user_message | TEXT | NOT NULL | User message")
el = body_after(el, "bot_response | TEXT | NOT NULL | Bot response")
el = body_after(el, "query_type | VARCHAR(50) | NOT NULL | Query category")
el = body_after(el, "created_at | DateTime | NOT NULL | Timestamp")

el = body_after(el, "Table 3.10 - Notification")
el = body_after(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body_after(el, "username | VARCHAR(100) | NOT NULL | Target user")
el = body_after(el, "message | TEXT | NOT NULL | Notification content")
el = body_after(el, "is_read | Boolean | DEFAULT FALSE | Read status")
el = body_after(el, "created_at | DateTime | NOT NULL | Timestamp")

print("3.5 Data Dictionary inserted")

# 3.6 ER Diagram - find last para before 4. Design heading again  
anchor = last_body_before('4. Design Document')
el = h2_after(anchor, "3.6 Entity Relationship Diagram")
el = body_after(el, "The Entity Relationship (ER) diagram illustrates the logical structure of the database with ten entities normalised up to Third Normal Form (3NF). The entities include roomtypedb, roomnamedb, staffdb, Registerdb, bookingdb, Totaldb, customercontactdb, ChatbotResponse, ChatbotConversation, and Notification. The diagram below shows the relationships between these tables.")
el = body_after(el, "Figure 3.3 - Entity Relationship Diagram")
print("3.6 ER Diagram inserted")

# 3.7 DFD
anchor = last_body_before('4. Design Document')
el = h2_after(anchor, "3.7 Data Flow Diagram")
el = body_after(el, "A Data Flow Diagram (DFD) represents the flow of data through the system.")
el = body_after(el, "Context Diagram (Level 0 DFD): The context diagram shows the entire system as a single process with external entities: Customer and Administrator. Key data flows include registration, login, search criteria, booking requests, payment details, chatbot queries from Customer to System; and room listings, booking confirmations, payment status, search results, chatbot responses from System to Customer.")
el = body_after(el, "Level 1 DFD: The system decomposes into eight processes: P1-User Management, P2-Room Management, P3-Search and Filtering, P4-Booking Engine, P5-Payment Processing, P6-Notification System, P7-Chatbot Engine, and P8-Admin Dashboard. Data stores include all ten database tables (D1-D10). External entities include Customer, Administrator, Razorpay Payment Gateway, and Gmail SMTP Server.")
el = body_after(el, "Figure 3.4 - Data Flow Diagram (Context and Level 1)")
print("3.7 DFD inserted")

print_headings("After Phase 2")
doc.save(SRC)
doc = Document(SRC)

# ============================================================
# PHASE 3: Add Security chapter + rename
# ============================================================
# Find last para before "8. Implementation Plan" heading
anchor = last_body_before('8. Implementation Plan')
print(f"\nPhase 3 anchor: {anchor.text.strip()[:60]}")

el = h1_after(anchor, "8. Implementation of Security for the Software Developed")
el = body_after(el, "Security is a critical aspect of any web application handling customer data and financial transactions. The system implements multiple security layers at application, database, and network levels.")
el = body_after(el, "8.1 User Authentication and Access Control: Role-based access with Customer and Administrator roles. Server-side sessions with encrypted data. Django auth with PBKDF2 password hashing for admin.")
el = body_after(el, "8.2 Password Security and OTP Verification: 5-digit OTP sent via email, stored server-side, single-use, validated with whitespace-stripped comparison.")
el = body_after(el, "8.3 CSRF Protection: Django CSRF tokens on all forms, validated on every POST/PUT/DELETE.")
el = body_after(el, "8.4 Payment Security: Razorpay PCI DSS Level 1. Payment on Razorpay's secure pages. Server-side signature verification.")
el = body_after(el, "8.5 SQL Injection Prevention: Django ORM parameterises all queries.")
el = body_after(el, "8.6 XSS Prevention: Django template engine auto-escapes all output.")
el = body_after(el, "8.7 Data Integrity: Three-level validation (client, server, model). Foreign key constraints.")
el = body_after(el, "8.8 Secure Email: SMTP with TLS encryption for notifications and OTPs.")

# Renumber
print("Renumbering...")
renames = [
    ("8. Implementation Plan", "9. Implementation Plan"),
    ("9. Limitations of the Project", "10. Limitations of the Project"),
    ("10. Future Application of the Project", "11. Future Application of the Project"),
    ("11. Conclusion", "12. Conclusion"),
    ("12. Bibliography", "13. Bibliography"),
]
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    for old, new in renames:
        if t == old or t.startswith(old):
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)

doc.save(SRC)
print("All phases complete!")

# Final verification
print_headings("FINAL STRUCTURE")
