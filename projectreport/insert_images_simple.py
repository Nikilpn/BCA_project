"""Insert diagram images simply by adding them to paragraphs."""
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import os

SRC = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx'
IMG_DIR = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport'

doc = Document(SRC)

def find_para(text_fragment):
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text and p.text.strip():
            return i, p
    return None, None

def add_image_para_after(anchor_para, img_path, width_inches=5.5):
    """Insert a new paragraph with an image right after anchor_para."""
    # Create a new paragraph element
    new_p_el = OxmlElement('w:p')
    
    # Create run
    run_el = OxmlElement('w:r')
    
    # Create drawing
    drawing_el = OxmlElement('w:drawing')
    
    # Add inline image using python-docx (via anchor_para)
    # Actually, python-docx's add_picture works on runs
    # Let's: add picture to anchor_para (temporarily), then move the drawing to our new paragraph
    
    # Easy way: just add to anchor_para then split
    run = anchor_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    
    # Now the picture run is at the end of anchor_para. 
    # Let's move it to a new paragraph after anchor_para.
    
    # Get the drawing element we just added (last run of anchor_para)
    last_run = anchor_para.runs[-1]
    drawing_elem = last_run._element.find(qn('w:drawing'))
    
    if drawing_elem is not None:
        # Move drawing to new paragraph
        # Remove it from the run
        run_elem = last_run._element
        run_elem.remove(drawing_elem)
        
        # If the run is now empty, keep it or clean it up
        # Create new paragraph
        new_p = OxmlElement('w:p')
        new_r = OxmlElement('w:r')
        new_r.append(drawing_elem)
        new_p.append(new_r)
        
        # Insert after anchor_para
        anchor_para._element.addnext(new_p)
        
        # If anchor_para's last run is now empty, we could remove it, but it's OK to keep
        return new_p
    
    return None

# Insert ER diagram
er_img = os.path.join(IMG_DIR, 'er_diagram.png')
if os.path.exists(er_img) and os.path.getsize(er_img) > 5000:
    _, p_er = find_para('Figure 3.3 - Entity Relationship Diagram')
    if p_er:
        add_image_para_after(p_er, er_img)
        print("ER image inserted")
    else:
        print("Figure 3.3 not found!")
else:
    print(f"ER image missing or invalid: size={os.path.getsize(er_img) if os.path.exists(er_img) else 0}")

# Insert DFD images (context + level 1)
dfd_ctx = os.path.join(IMG_DIR, 'dfd_context.png')
dfd_l1 = os.path.join(IMG_DIR, 'dfd_level1.png')

if os.path.exists(dfd_ctx) and os.path.getsize(dfd_ctx) > 5000:
    _, p_dfd = find_para('Figure 3.4')
    if p_dfd:
        add_image_para_after(p_dfd, dfd_ctx)
        print("DFD Context image inserted")
        
        # Insert Level 1 after context
        if os.path.exists(dfd_l1) and os.path.getsize(dfd_l1) > 5000:
            # Re-find figure 3.4 text (now shift)
            _, p_dfd2 = find_para('Figure 3.4')
            if p_dfd2:
                # Need to insert after the DFD context image paragraph, 
                # Find the next sibling of DFD context paragraph
                # p_dfd2 is the figure caption, its next sibling is the DFD context image
                next_sib = p_dfd2._element.getnext()
                if next_sib is not None:
                    # Insert after next_sib
                    new_p = OxmlElement('w:p')
                    run = doc.add_paragraph().add_run()
                    run.add_picture(dfd_l1, width=Inches(5.5))
                    # Similar movement
                    last_run = run._element
                    # Actually, this is getting complex. Let's use a simpler approach.
                    pass

# Actually let me take a completely different approach using add_picture
# on a fresh paragraph

print("Using simpler approach...")

# Re-open (images were added via add_run, so we need clean state)
doc = Document(SRC)

# Reset: save without images first
doc.save(SRC)
doc = Document(SRC)

def add_image_separate_paragraph(anchor_para, img_path, width_inches=5.5):
    """Add an image in its own paragraph after anchor paragraph."""
    # Add a new paragraph with just an empty run
    # Use XML to insert an empty paragraph
    ref = anchor_para._element
    
    # Create image paragraph: add picture to current doc to get the relationship
    # Then extract what we need
    rId = doc.part.relate_to(img_path, 
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image',
        is_external=False)
    
    # Create paragraph
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    
    # Add drawing
    drawing = OxmlElement('w:drawing')
    
    # wp:inline
    wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    inline = OxmlElement('{%s}inline' % wp_ns)
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '0')
    inline.set('distR', '0')
    
    extent = OxmlElement('{%s}extent' % wp_ns)
    w_emu = str(int(width_inches * 914400))
    h_emu = str(int(width_inches * 0.65 * 914400))  # 65% aspect ratio
    extent.set('cx', w_emu)
    extent.set('cy', h_emu)
    inline.append(extent)
    
    effectExtent = OxmlElement('{%s}effectExtent' % wp_ns)
    effectExtent.set('l', '0')
    effectExtent.set('t', '0')
    effectExtent.set('r', '0')
    effectExtent.set('b', '0')
    inline.append(effectExtent)
    
    docPr = OxmlElement('{%s}docPr' % wp_ns)
    docPr.set('id', '1')
    docPr.set('name', 'Image')
    inline.append(docPr)
    
    # Graphic
    a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    graphic = OxmlElement('{%s}graphic' % a_ns)
    
    graphicData = OxmlElement('{%s}graphicData' % a_ns)
    graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')
    
    # pic:pic
    pic_ns = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    pic = OxmlElement('{%s}pic' % pic_ns)
    
    nvPicPr = OxmlElement('{%s}nvPicPr' % pic_ns)
    cNvPr = OxmlElement('{%s}cNvPr' % pic_ns)
    cNvPr.set('id', '0')
    cNvPr.set('name', 'Picture')
    nvPicPr.append(cNvPr)
    cNvPicPr = OxmlElement('{%s}cNvPicPr' % pic_ns)
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)
    
    blipFill = OxmlElement('{%s}blipFill' % pic_ns)
    blip = OxmlElement('{%s}blip' % a_ns)
    blip.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', rId)
    blipFill.append(blip)
    srcRect = OxmlElement('{%s}srcRect' % a_ns)
    blipFill.append(srcRect)
    stretch = OxmlElement('{%s}stretch' % a_ns)
    fillRect = OxmlElement('{%s}fillRect' % a_ns)
    stretch.append(fillRect)
    blipFill.append(stretch)
    pic.append(blipFill)
    
    spPr = OxmlElement('{%s}spPr' % pic_ns)
    xfrm = OxmlElement('{%s}xfrm' % a_ns)
    off = OxmlElement('{%s}off' % a_ns)
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext = OxmlElement('{%s}ext' % a_ns)
    ext.set('cx', w_emu)
    ext.set('cy', h_emu)
    xfrm.append(ext)
    spPr.append(xfrm)
    prstGeom = OxmlElement('{%s}prstGeom' % a_ns)
    prstGeom.set('prst', 'rect')
    spPr.append(prstGeom)
    noFill = OxmlElement('{%s}noFill' % a_ns)
    spPr.append(noFill)
    pic.append(spPr)
    
    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)
    drawing.append(inline)
    new_r.append(drawing)
    new_p.append(new_r)
    
    ref.addnext(new_p)
    return new_p

# Insert ER
_, p_er = find_para('Figure 3.3 - Entity Relationship Diagram')
if p_er and os.path.exists(er_img) and os.path.getsize(er_img) > 5000:
    add_image_separate_paragraph(p_er, er_img)
    print("ER diagram image inserted")

# Insert DFD context
_, p_dfd = find_para('Figure 3.4')
if p_dfd and os.path.exists(dfd_ctx) and os.path.getsize(dfd_ctx) > 5000:
    add_image_separate_paragraph(p_dfd, dfd_ctx)
    print("DFD Context image inserted")
    
    # Insert DFD Level 1 after Context
    if os.path.exists(dfd_l1) and os.path.getsize(dfd_l1) > 5000:
        add_image_separate_paragraph(p_dfd, dfd_l1)
        print("DFD Level 1 image inserted")

doc.save(SRC)
print("All images inserted and saved!")
