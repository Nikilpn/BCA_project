"""Build Data Dictionary tables directly in XML at correct position."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx'
doc = Document(SRC)

def make_table_xml(rows_data, col_widths_cm):
    """Create a w:tbl XML element with 4 columns and header+data rows."""
    W = 'w:'  # prefix
    
    tbl = OxmlElement(W + 'tbl')
    
    # Table properties
    tblPr = OxmlElement(W + 'tblPr')
    tblW = OxmlElement(W + 'tblW')
    tblW.set(qn(W + 'w'), '9000')
    tblW.set(qn(W + 'type'), 'dxa')
    tblPr.append(tblW)
    tblBorders = OxmlElement(W + 'tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(W + edge)
        border.set(qn(W + 'val'), 'single')
        border.set(qn(W + 'sz'), '4')
        border.set(qn(W + 'color'), '333333')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tblLook = OxmlElement(W + 'tblLook')
    tblLook.set(qn(W + 'val'), '04A0')
    tblPr.append(tblLook)
    tbl.append(tblPr)
    
    def make_cell(text, bold=False, bg_color=None):
        cell = OxmlElement(W + 'tc')
        tcPr = OxmlElement(W + 'tcPr')
        tcW = OxmlElement(W + 'tcW')
        tcW.set(qn(W + 'w'), '2000')
        tcW.set(qn(W + 'type'), 'dxa')
        tcPr.append(tcW)
        if bg_color:
            shading = OxmlElement(W + 'shd')
            shading.set(qn(W + 'fill'), bg_color)
            shading.set(qn(W + 'val'), 'clear')
            tcPr.append(shading)
        cell.append(tcPr)
        p = OxmlElement(W + 'p')
        pPr = OxmlElement(W + 'pPr')
        pStyle = OxmlElement(W + 'pStyle')
        pStyle.set(qn(W + 'val'), 'Normal')
        pPr.append(pStyle)
        p.append(pPr)
        r = OxmlElement(W + 'r')
        rPr = OxmlElement(W + 'rPr')
        rFonts = OxmlElement(W + 'rFonts')
        rFonts.set(qn(W + 'ascii'), 'Calibri')
        rFonts.set(qn(W + 'hAnsi'), 'Calibri')
        rPr.append(rFonts)
        sz = OxmlElement(W + 'sz')
        sz.set(qn(W + 'val'), '18')  # 9pt = 18 half-points
        rPr.append(sz)
        if bold:
            b = OxmlElement(W + 'b')
            rPr.append(b)
        r.append(rPr)
        t = OxmlElement(W + 't')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
        cell.append(p)
        return cell
    
    # Header row
    header_labels = ['Field Name', 'Data Type', 'Constraints', 'Description']
    tr_header = OxmlElement(W + 'tr')
    for h in header_labels:
        tr_header.append(make_cell(h, bold=True, bg_color='2F5496'))
    tbl.append(tr_header)
    
    # Data rows
    for ri, row in enumerate(rows_data):
        tr = OxmlElement(W + 'tr')
        bg = 'D6E4F0' if ri % 2 == 0 else None
        for cell_val in row:
            tr.append(make_cell(cell_val, bg_color=bg))
        tbl.append(tr)
    
    return tbl

# Find Data Dictionary boundaries and collect data
dd_start = None
dd_end = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == '3.5 Data Dictionary' and 'Heading' in p.style.name:
        dd_start = i
    if dd_start is not None and i > dd_start:
        if '3.6 Entity Relationship' in t and 'Heading' in p.style.name:
            dd_end = i
            break

print(f"DD: P{dd_start} to P{dd_end-1}")

# Collect data
dd_data = {}
current_table = None
intro_para = None

for i in range(dd_start + 1, dd_end):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if i == dd_start + 1:
        intro_para = p
        continue
    if t.startswith('Table 3.'):
        current_table = t
        dd_data[current_table] = []
    elif current_table and '|' in t:
        parts = [x.strip() for x in t.split('|')]
        # Ensure 4 parts
        while len(parts) < 4:
            parts.append('')
        dd_data[current_table].append(parts[:4])

print(f"Tables: {[k[:30] for k in dd_data.keys()]}")
print(f"Total rows: {sum(len(v) for v in dd_data.values())}")

# Remove all DD content paragraphs (keep heading + intro)
elements_to_remove = []
for i in range(dd_start + 2, dd_end):
    elements_to_remove.append(doc.paragraphs[i]._element)

BODY = doc.element.body
for e in elements_to_remove:
    try:
        BODY.remove(e)
    except:
        pass

print(f"Removed {len(elements_to_remove)} paragraphs")

# Now insert table_name paragraphs + tables after intro_para
insert_point = intro_para._element

col_widths = [3, 3, 4, 7]  # in cm

for table_name, rows in dd_data.items():
    # Insert table name paragraph
    name_p = OxmlElement('w:p')
    name_r = OxmlElement('w:r')
    name_t = OxmlElement('w:t')
    name_t.text = table_name
    name_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    name_r.append(name_t)
    name_p.append(name_r)
    insert_point.addnext(name_p)
    insert_point = name_p
    
    # Insert table XML
    tbl_xml = make_table_xml(rows, col_widths)
    insert_point.addnext(tbl_xml)
    insert_point = tbl_xml

print("All tables inserted!")

doc.save(SRC)
print("Saved!")

# Quick verify
doc2 = Document(SRC)
tables = doc2.tables
print(f"\nTotal tables in doc: {len(tables)}")
# Find DD tables
count = 0
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if 'Table 3.' in t:
        # Check next sibling for table
        next_elem = p._element.getnext()
        if next_elem is not None and next_elem.tag == qn('w:tbl'):
            count += 1
            # Find this table in doc.tables
            for ti, table in enumerate(doc2.tables):
                first_cell = table.rows[0].cells[0].text.strip() if table.rows else ''
                if first_cell == 'Field Name':
                    print(f"  ✓ {t} -> {len(table.rows)} rows (header + {len(table.rows)-1} data)")
                    break
print(f"Verified {count}/10 DD tables with correct format")
