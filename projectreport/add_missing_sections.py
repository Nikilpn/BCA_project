"""Add ONLY missing sections to nikhilproject241.docx.
Works bottom-up so paragraph indices remain valid.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

doc = Document('/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx')

def add_after(para_or_el, text, style_name=None):
    """Insert a new paragraph after the given element. Returns the new XML element."""
    ref = para_or_el._element if hasattr(para_or_el, '_element') else para_or_el
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    new_p.append(r)
    if style_name:
        pPr = ref.find(qn('w:pPr'))
        if pPr is not None:
            new_pPr = deepcopy(pPr)
            pStyle = new_pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                sid = doc.styles[style_name].element.get(qn('w:styleId'))
                pStyle.set(qn('w:val'), sid)
            new_p.insert(0, new_pPr)
    ref.addnext(new_p)
    return new_p

def body(el, text):
    return add_after(el, text, 'Body Text')

def h1(el, text):
    return add_after(el, text, 'Heading 1.chapter')

def h2(el, text):
    return add_after(el, text, 'Heading 2.section')

def h3(el, text):
    return add_after(el, text, 'Heading 3.subsec')

def find_last_para_before(heading_prefix):
    """Find the last non-empty paragraph before the heading starting with prefix."""
    idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(heading_prefix) and 'Heading' in p.style.name:
            idx = i
            break
    if idx is None:
        return None, None
    anchor = idx - 1
    while anchor >= 0 and not doc.paragraphs[anchor].text.strip():
        anchor -= 1
    return doc.paragraphs[anchor], doc.paragraphs[idx]


# ============================================================
# STEP 1: Add Implementation of Security chapter (after Ch 7, before Ch 8)
# Ch 7 ends at last para before "8. Implementation Plan" heading
# ============================================================
anchor_el, _ = find_last_para_before('8. Implementation Plan')
print(f"Security chapter: inserting after P that starts: {anchor_el.text.strip()[:80]}")

el = h1(anchor_el, "8. Implementation of Security for the Software Developed")
el = body(el,
    "Security is a critical aspect of any web application that handles customer data and financial transactions. "
    "The Hotel Room Booking Website with Assistant Chatbot implements multiple layers of security at the application, "
    "database, and network levels to protect user data, prevent unauthorised access, and ensure safe payment processing. "
    "This chapter documents the security measures implemented in the system.")

el = body(el, "8.1 User Authentication and Access Control")
el = body(el,
    "The system implements a role-based access control (RBAC) model with two distinct user roles: Customers and Administrators. "
    "Customer authentication is handled through a registration and login system with server-side session management. "
    "Upon successful login, the username is stored in the Django session object, which is encrypted and signed "
    "using Django's session middleware. Session data is stored server-side in the database-backed session store, "
    "preventing client-side tampering. Session expiration is enforced and sessions are destroyed on logout.")
el = body(el,
    "Administrator access is secured through Django's built-in authentication system which uses "
    "PBKDF2 with SHA-256 hashing for password storage. The admin interface enforces permission-based access control, "
    "restricting administrative functions to authenticated staff users only.")

el = body(el, "8.2 Password Security and OTP Verification")
el = body(el,
    "The password reset feature implements multi-factor verification using a One-Time Password (OTP) sent via "
    "email. The OTP is a randomly generated 5-digit number stored in the server-side session, ensuring it cannot "
    "be accessed or modified by the client. On submission, both the user input and stored OTP are compared after "
    "whitespace stripping. The OTP is single-use and removed from the session after verification.")

el = body(el, "8.3 CSRF Protection")
el = body(el,
    "All HTML forms include Django's CSRF token, validated by CsrfViewMiddleware on every POST, PUT, and DELETE "
    "request. Requests with missing or invalid tokens are rejected with a 403 Forbidden response, preventing "
    "cross-site request forgery attacks.")

el = body(el, "8.4 Payment Security via Razorpay")
el = body(el,
    "Payment processing is handled by Razorpay (PCI DSS Level 1 compliant). Sensitive payment information is "
    "collected on Razorpay's secure checkout page, never on the application's servers. Payment verification "
    "uses server-side signature validation. Razorpay API keys are stored server-side in Django settings.")

el = body(el, "8.5 SQL Injection Prevention")
el = body(el,
    "Django's ORM is used for all database queries, automatically parameterising inputs and treating them as "
    "query parameters rather than executable SQL. No raw SQL queries are executed in the application.")

el = body(el, "8.6 XSS Prevention")
el = body(el,
    "Django's template engine auto-escapes all variable output, converting HTML special characters to their "
    "entity equivalents. Explicit safe filters are used only for trusted static content.")

el = body(el, "8.7 Data Integrity and Validation")
el = body(el,
    "Input validation is implemented at three levels: HTML5 client-side validation, Django view-level server-side "
    "validation (including date range checks and email format validation), and model-level type enforcement. "
    "Referential integrity is maintained through foreign key constraints with appropriate on_delete behaviours.")

el = body(el, "8.8 Secure Email Communication")
el = body(el,
    "Email notifications use Django's SMTP backend with TLS encryption for booking confirmations and OTP delivery. "
    "Email content does not include passwords or sensitive payment information.")

el = body(el, "In summary, the Hotel Room Booking System implements comprehensive security measures across all "
    "layers, providing essential protection against CSRF, XSS, SQL injection, and payment data exposure.")


# ============================================================
# STEP 2: Add Data Flow Diagram (3.7), ER Diagram (3.6), Data Dictionary (3.5)
# Insert before Ch 4 heading, work bottom-up
# ============================================================

# --- 2a. Data Flow Diagram (3.7) ---
anchor_el, _ = find_last_para_before('4. Design Document')
print(f"DFD section: inserting after P that starts: {anchor_el.text.strip()[:80]}")

el = h2(anchor_el, "3.7 Data Flow Diagram")
el = body(el,
    "A Data Flow Diagram (DFD) represents the flow of data through the Hotel Room Booking System, showing "
    "how information moves between external entities, processes, and data stores at different levels of abstraction.")
el = body(el, "Context Diagram (Level 0 DFD):")
el = body(el,
    "The context diagram shows the entire system as a single process with two external entities: Customer and "
    "Administrator. Data flows from Customer to System include: Registration Details, Login Credentials, Room "
    "Search Criteria, Booking Request, Payment Details, Contact Enquiry, and Chatbot Queries. Data flows from "
    "System to Customer include: Room Listings, Booking Confirmation, Payment Status, Search Results, Chatbot "
    "Responses, Email Notifications, and PDF Receipt. Data flows from Administrator to System include: Admin "
    "Login, Room Type Data, Room Data, Staff Data. Data flows from System to Administrator include: Booking "
    "Records, Payment Records, Contact Messages, and Dashboard Summary.")
el = body(el, "Level 1 DFD - Main Processes:")
el = body(el,
    "The system is decomposed into eight processes: Process 1 - User Management (Registration, Login, Password "
    "Reset); Process 2 - Room Management (CRUD operations); Process 3 - Search and Filtering (AJAX-based with "
    "date availability); Process 4 - Booking Engine (Date overlap detection, cart management); Process 5 - "
    "Payment Processing (Razorpay integration); Process 6 - Notification System (Email, OTP); Process 7 - "
    "Chatbot Engine (Query classification, response generation); Process 8 - Admin Dashboard.")
el = body(el,
    "Data Stores: D1 - roomtypedb, D2 - roomnamedb, D3 - staffdb, D4 - Registerdb, D5 - bookingdb, "
    "D6 - Totaldb, D7 - customercontactdb, D8 - ChatbotResponse, D9 - ChatbotConversation, D10 - Notification.")
el = body(el,
    "External Entities: Customer, Administrator, Razorpay Payment Gateway, Gmail SMTP Server.")
el = body(el, "Figure 3.2 - Data Flow Diagram (Context and Level 1)")


# --- 2b. Entity Relationship Diagram (3.6) ---
anchor_el, _ = find_last_para_before('4. Design Document')
print(f"ER section: inserting after P that starts: {anchor_el.text.strip()[:80]}")

el = h2(anchor_el, "3.6 Entity Relationship Diagram")
el = body(el,
    "The Entity Relationship (ER) diagram illustrates the logical structure of the database with ten entities "
    "normalised up to Third Normal Form (3NF).")

el = body(el, "Entities: roomtypedb (Room Category), roomnamedb (Room), staffdb (Staff Profile), Registerdb "
    "(Customer Registration), bookingdb (Booking), Totaldb (Payment), customercontactdb (Contact Enquiry), "
    "ChatbotResponse (Training Data), ChatbotConversation (Conversation Log), Notification.")

el = body(el, "Relationships:")
el = body(el,
    "1. Room Type to Room (One-to-Many): One room type can have many rooms. Foreign key ROOMTYPE in roomnamedb "
    "references roomtypedb.id.")
el = body(el,
    "2. Customer to Booking (One-to-Many): One customer can make multiple bookings. Foreign key CUSTOMER in "
    "bookingdb references Registerdb.id with SET_NULL on delete.")
el = body(el,
    "3. Room to Booking (One-to-Many): One room can appear in multiple bookings. Foreign key SELECTROOM in "
    "bookingdb references roomnamedb.id with CASCADE on delete.")
el = body(el,
    "4. Booking to Payment (One-to-One): Each booking has at most one payment. Foreign key BOOKING in Totaldb "
    "references bookingdb.id with CASCADE on delete.")
el = body(el, "Figure 3.3 - Entity Relationship Diagram")


# --- 2c. Data Dictionary (3.5) ---
anchor_el, _ = find_last_para_before('4. Design Document')
print(f"Data Dictionary: inserting after P that starts: {anchor_el.text.strip()[:80]}")

el = h2(anchor_el, "3.5 Data Dictionary")
el = body(el,
    "The data dictionary defines the structure, data types, constraints, and descriptions for all database tables.")

# roomtypedb
el = body(el, "Table 3.1 - roomtypedb (Room Category)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "ROOMTYPE | VARCHAR(100) | NOT NULL | Category name (Luxury, Deluxe, Standard, Budget)")
el = body(el, "ROOMTYPEIMAGE | VARCHAR(255) | NULLABLE | Image file path")
el = body(el, "DESCRIPTION | TEXT | NULLABLE | Category description")

# roomnamedb
el = body(el, "Table 3.2 - roomnamedb (Room)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "ROOMTYPE | Integer | FK (roomtypedb.id) | Room category reference")
el = body(el, "ROOMNAME | VARCHAR(100) | NOT NULL | Room name/number")
el = body(el, "ROOMIMAGE | VARCHAR(255) | NULLABLE | Image file path")
el = body(el, "ROOMPRICE | Integer | NOT NULL | Price per night (INR)")
el = body(el, "ROOMDESCRIPTION | TEXT | NULLABLE | Room features")

# staffdb
el = body(el, "Table 3.3 - staffdb (Staff Profile)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "STAFFNAME | VARCHAR(100) | NOT NULL | Staff full name")
el = body(el, "DESIGNATION | VARCHAR(100) | NOT NULL | Job title")
el = body(el, "FACEBOOKURL | VARCHAR(255) | NULLABLE | Facebook URL")
el = body(el, "INSTAGRAMURL | VARCHAR(255) | NULLABLE | Instagram URL")
el = body(el, "STAFFIMAGE | VARCHAR(255) | NULLABLE | Photograph path")

# Registerdb
el = body(el, "Table 3.4 - Registerdb (Customer Registration)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "USERNAME | VARCHAR(100) | UNIQUE, NOT NULL | Login username")
el = body(el, "PASSWORD | VARCHAR(100) | NOT NULL | Password")
el = body(el, "CONFIRMPASSWORD | VARCHAR(100) | NOT NULL | Password confirmation")
el = body(el, "EMAIL | VARCHAR(100) | NOT NULL | Email address")

# bookingdb
el = body(el, "Table 3.5 - bookingdb (Booking Record)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "CUSTOMER | Integer | FK (Registerdb.id), SET_NULL | Customer reference")
el = body(el, "CUSTOMERNAME | VARCHAR(100) | NOT NULL | Customer name")
el = body(el, "CONTACTEMAIL | VARCHAR(100) | NOT NULL | Customer email")
el = body(el, "CHECKIN | Date | NOT NULL | Check-in date")
el = body(el, "CHECKOUT | Date | NOT NULL | Check-out date")
el = body(el, "TOTALADULTS | Integer | NOT NULL, DEFAULT 1 | Adult count")
el = body(el, "TOTALCHILDS | Integer | NOT NULL, DEFAULT 0 | Child count")
el = body(el, "SELECTROOM | Integer | FK (roomnamedb.id), CASCADE | Room reference")
el = body(el, "SPECIALREQUEST | TEXT | NULLABLE | Special requests")
el = body(el, "TOTALPRICE | Integer | NOT NULL | Total price (INR)")

# Totaldb
el = body(el, "Table 3.6 - Totaldb (Payment Transaction)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "BOOKING | Integer | FK (bookingdb.id), CASCADE | Booking reference")
el = body(el, "CUSTOMERNAME | VARCHAR(100) | NOT NULL | Customer name")
el = body(el, "MOBILE | VARCHAR(15) | NOT NULL | Mobile number")
el = body(el, "TOTALPRICE | Integer | NOT NULL | Amount paid (INR)")

# customercontactdb
el = body(el, "Table 3.7 - customercontactdb (Contact Enquiry)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "CONTACTNAME | VARCHAR(100) | NOT NULL | Sender name")
el = body(el, "CONTACTNUMBER | VARCHAR(15) | NULLABLE | Phone number")
el = body(el, "CONTACTEMAIL | VARCHAR(100) | NOT NULL | Sender email")
el = body(el, "CONTACTSUBJECT | VARCHAR(200) | NOT NULL | Subject")
el = body(el, "CONTACTMESSAGE | TEXT | NOT NULL | Message content")

# ChatbotResponse
el = body(el, "Table 3.8 - ChatbotResponse")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "keywords | TEXT | NOT NULL | Keywords for pattern matching")
el = body(el, "response | TEXT | NOT NULL | Bot response text")
el = body(el, "query_type | VARCHAR(50) | NOT NULL | Query category")

# ChatbotConversation
el = body(el, "Table 3.9 - ChatbotConversation")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "username | VARCHAR(100) | NOT NULL | Participant username")
el = body(el, "user_message | TEXT | NOT NULL | User message")
el = body(el, "bot_response | TEXT | NOT NULL | Bot response")
el = body(el, "query_type | VARCHAR(50) | NOT NULL | Query category")
el = body(el, "created_at | DateTime | NOT NULL | Timestamp")

# Notification
el = body(el, "Table 3.10 - Notification")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "username | VARCHAR(100) | NOT NULL | Target user")
el = body(el, "message | TEXT | NOT NULL | Notification content")
el = body(el, "is_read | Boolean | DEFAULT FALSE | Read status")
el = body(el, "created_at | DateTime | NOT NULL | Timestamp")


# Save
output_path = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx'
doc.save(output_path)
print(f"\nSaved: {output_path}")
print("ALL MISSING SECTIONS ADDED SUCCESSFULLY!")
