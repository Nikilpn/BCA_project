"""Insert diagram images using python-docx native API."""
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from zipfile import ZipFile
from lxml import etree
import os, shutil

SRC = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx'
IMG_DIR = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport'
TMP = SRC + '.bak2.docx'

# Make a temp copy, then rebuild with images added via docx API
doc = Document(SRC)

def find_para(text_fragment):
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text and p.text.strip():
            return i, p
    return None, None

def add_image_via_paragraph(anchor_para, img_path, width_inches=5.5):
    """Add image by inserting a new paragraph with add_picture."""
    # Need to: 
    # 1. Create new paragraph in the document body
    # 2. Add picture to it  
    # 3. Move the paragraph to right after anchor_para
    # 
    # add_picture adds to runs of a paragraph. The hard part is moving it.
    # Easiest: temporarily add to a new paragraph at end of doc, 
    # then move the XML element.
    
    # Create temp paragraph at end
    temp_para = doc.add_paragraph()
    run = temp_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    temp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Get the temp paragraph's element
    temp_elem = temp_para._element
    
    # Move it right after anchor_para
    anchor_para._element.addnext(temp_elem)
    
    return temp_elem

# ER diagram
er_path = os.path.join(IMG_DIR, 'er_diagram.png')
_, p_er = find_para('Figure 3.3')
if p_er and os.path.exists(er_path) and os.path.getsize(er_path) > 5000:
    add_image_via_paragraph(p_er, er_path)
    print("✓ ER diagram inserted")

# DFD Context
dfd_ctx = os.path.join(IMG_DIR, 'dfd_context.png')
_, p_dfd = find_para('Figure 3.4')
if p_dfd and os.path.exists(dfd_ctx) and os.path.getsize(dfd_ctx) > 5000:
    add_image_via_paragraph(p_dfd, dfd_ctx)
    print("✓ DFD Context inserted")

# DFD Level 1
dfd_l1 = os.path.join(IMG_DIR, 'dfd_level1.png')
if os.path.exists(dfd_l1) and os.path.getsize(dfd_l1) > 5000:
    # Insert after the context image para (which is after Figure 3.4 para)
    add_image_via_paragraph(p_dfd, dfd_l1)
    print("✓ DFD Level 1 inserted")

doc.save(SRC)
print("Saved!")
