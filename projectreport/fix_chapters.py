import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

doc = docx.Document('/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject24.docx')

def insert_after(para, text, style_name=None):
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    new_p.append(r)
    if style_name:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            new_pPr = deepcopy(pPr)
            pStyle = new_pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                sid = doc.styles[style_name].element.get(qn('w:styleId'))
                pStyle.set(qn('w:val'), sid)
            new_p.insert(0, new_pPr)
    para._element.addnext(new_p)
    return new_p

def body(para, text):
    return insert_after(para, text, 'Body Text')

def h1(para, text):
    return insert_after(para, text, 'Heading 1.chapter')

def h2(para, text):
    return insert_after(para, text, 'Heading 2.section')

# First, let's find ALL paragraphs and identify the actual structure
print("=== CURRENT STATE (all heading-level paragraphs with index) ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        s = p.style.name
        if 'Heading' in s:
            print(f"  P{i}: [{s}] {p.text.strip()[:80]}")

# Get all chapter headings with their XML elements
chapters = []
for i, p in enumerate(doc.paragraphs):
    if 'Heading 1' in p.style.name:
        chapters.append((i, p))

print(f"\nFound {len(chapters)} chapter headings")

# Now let's figure out what each chapter should be numbered
# Ordered list of chapters by their appearance:
# Certificate of Originality (no number)
# Acknowledgement (no number)
# Abstract (no number)
# Table of Contents (no number)
# 1. Introduction
# 2. Tools and Environment Used
# 3. Requirement Specification
# 4. Design Document
# 5. System Architecture and Design Patterns (NEW)
# 6. Program Code
# 7. Testing
# 8. Input Output Screens
# 9. Implementation Plan
# 10. Limitations of the Project
# 11. Future Application of the Project
# 12. Conclusion
# 13. Software Engineering and SDLC Approach
# 14. Security Analysis and Implementation (NEW)
# 15. Bibliography
# THANK YOU

