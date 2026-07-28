import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import re

doc = docx.Document('/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx')

renames = [
    ("8. Implementation Plan", "9. Implementation Plan"),
    ("9. Limitations", "10. Limitations"),
    ("10. Future Application", "11. Future Application"),
    ("11. Conclusion", "12. Conclusion"),
    ("12. Bibliography", "13. Bibliography"),
]

for p in doc.paragraphs:
    t = p.text.strip()
    for old, new in renames:
        if t == old:
            # Replace the text in the paragraph
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    # Change heading style to keep formatting
            print(f"Renamed: '{old}' -> '{new}'")

# Also renumber Table of Contents entries if they exist
for p in doc.paragraphs:
    t = p.text.strip()
    for old, new in renames:
        if t.startswith(old):
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
            break

output_path = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx'
doc.save(output_path)
print(f"\nDocument saved: {output_path}")
