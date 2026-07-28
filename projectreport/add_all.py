"""Fix 3.4/3.3 order, add 3.5/3.6/3.7 in correct positions, add Security chapter."""
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

SRC = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx'
doc = Document(SRC)
BODY_EL = doc.element.body

def add_after(para, text, style_name=None):
    ref = para._element
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    new_p.append(r)
    if style_name:
        pPr = ref.find(qn('w:pPr'))
        if pPr is not None:
            new_pPr = deepcopy(pPr)
            pStyle = new_pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                sid = doc.styles[style_name].element.get(qn('w:styleId'))
                pStyle.set(qn('w:val'), sid)
            new_p.insert(0, new_pPr)
    ref.addnext(new_p)
    return new_p

def h2(p, t): return add_after(p, t, 'Heading 2.section')
def h1(p, t): return add_after(p, t, 'Heading 1.chapter')
def body(p, t): return add_after(p, t, 'Body Text')

def find_start(heading_text):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading_text and 'Heading' in p.style.name:
            return i, p
    return None, None

def section_range(start_idx):
    """Return (start_idx, end_idx_exclusive) for a section."""
    end = len(doc.paragraphs)
    for i in range(start_idx + 1, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        s = doc.paragraphs[i].style.name
        if 'Heading' in s and (t.startswith('3.') or t.startswith('4.')):
            end = i
            break
    return start_idx, end

# ============================================================
# STEP 1: Fix 3.4/3.3 ordering by moving 3.4 section after 3.3
# Current: ... → 3.4 (P179-184) → 3.3 (P185-187) → 4. Design (P188)
# Target:  ... → 3.3 → 3.4 → 4. Design
# ============================================================

# Get indices
idx_34, p_34 = find_start('3.4 Software and Hardware Requirements')
idx_33, p_33 = find_start('3.3 Use Case Diagram')

print(f"3.4 at P{idx_34}, 3.3 at P{idx_33}")

# Get 3.4 and 3.3 ranges
s34_start, s34_end = section_range(idx_34)
s33_start, s33_end = section_range(idx_33)

print(f"3.4 range: P{s34_start}-{s34_end-1}")
print(f"3.3 range: P{s33_start}-{s33_end-1}")

# Collect XML elements for 3.4 section
el_34 = [doc.paragraphs[i]._element for i in range(s34_start, s34_end)]

# Reference element: the last element of 3.3 section (to insert 3.4 after 3.3)
last_33_elem = doc.paragraphs[s33_end-1]._element

# Remove 3.4 elements from their current positions
for e in el_34:
    BODY_EL.remove(e)

# Insert 3.4 elements after last_33_elem
for e in el_34:
    last_33_elem.addnext(e)
    last_33_elem = e

print("3.4 moved after 3.3")

# ============================================================
# STEP 2: Insert 3.5 Data Dictionary after 3.4's last para
# ============================================================
# Find the last element of 3.4 (which is now after 3.3)
idx_34, p_34 = find_start('3.4 Software and Hardware Requirements')
s34_start, s34_end = section_range(idx_34)
anchor_34 = doc.paragraphs[s34_end - 1]
print(f"Last element of 3.4: P{s34_end-1}: {anchor_34.text.strip()[:60]}")

el = h2(anchor_34, "3.5 Data Dictionary")
el = body(el, "The data dictionary defines the structure, data types, constraints, and descriptions for all database tables.")

def dd_table(anchor, table_name, rows):
    """Insert a data dictionary table."""
    el = body(anchor, table_name)
    for r in rows:
        el = body(el, r)
    return el

# roomtypedb
el = body(el, "Table 3.1 - roomtypedb (Room Category)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "ROOMTYPE | VARCHAR(100) | NOT NULL | Category name (Luxury, Deluxe, Standard, Budget)")
el = body(el, "ROOMTYPEIMAGE | VARCHAR(255) | NULLABLE | Image file path")
el = body(el, "DESCRIPTION | TEXT | NULLABLE | Category description")

# roomnamedb
el = body(el, "Table 3.2 - roomnamedb (Room)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "ROOMTYPE | Integer | FK (roomtypedb.id) | Room category reference")
el = body(el, "ROOMNAME | VARCHAR(100) | NOT NULL | Room name/number")
el = body(el, "ROOMIMAGE | VARCHAR(255) | NULLABLE | Image file path")
el = body(el, "ROOMPRICE | Integer | NOT NULL | Price per night (INR)")
el = body(el, "ROOMDESCRIPTION | TEXT | NULLABLE | Room features")

# staffdb
el = body(el, "Table 3.3 - staffdb (Staff Profile)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "STAFFNAME | VARCHAR(100) | NOT NULL | Staff full name")
el = body(el, "DESIGNATION | VARCHAR(100) | NOT NULL | Job title")
el = body(el, "FACEBOOKURL | VARCHAR(255) | NULLABLE | Facebook URL")
el = body(el, "INSTAGRAMURL | VARCHAR(255) | NULLABLE | Instagram URL")
el = body(el, "STAFFIMAGE | VARCHAR(255) | NULLABLE | Photograph path")

# Registerdb
el = body(el, "Table 3.4 - Registerdb (Customer Registration)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "USERNAME | VARCHAR(100) | UNIQUE, NOT NULL | Login username")
el = body(el, "PASSWORD | VARCHAR(100) | NOT NULL | Password")
el = body(el, "CONFIRMPASSWORD | VARCHAR(100) | NOT NULL | Password confirmation")
el = body(el, "EMAIL | VARCHAR(100) | NOT NULL | Email address")

# bookingdb
el = body(el, "Table 3.5 - bookingdb (Booking Record)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "CUSTOMER | Integer | FK (Registerdb.id), SET_NULL | Customer reference")
el = body(el, "CUSTOMERNAME | VARCHAR(100) | NOT NULL | Customer name")
el = body(el, "CONTACTEMAIL | VARCHAR(100) | NOT NULL | Customer email")
el = body(el, "CHECKIN | Date | NOT NULL | Check-in date")
el = body(el, "CHECKOUT | Date | NOT NULL | Check-out date")
el = body(el, "TOTALADULTS | Integer | NOT NULL, DEFAULT 1 | Adult count")
el = body(el, "TOTALCHILDS | Integer | NOT NULL, DEFAULT 0 | Child count")
el = body(el, "SELECTROOM | Integer | FK (roomnamedb.id), CASCADE | Room reference")
el = body(el, "SPECIALREQUEST | TEXT | NULLABLE | Special requests")
el = body(el, "TOTALPRICE | Integer | NOT NULL | Total price (INR)")

# Totaldb
el = body(el, "Table 3.6 - Totaldb (Payment Transaction)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "BOOKING | Integer | FK (bookingdb.id), CASCADE | Booking reference")
el = body(el, "CUSTOMERNAME | VARCHAR(100) | NOT NULL | Customer name")
el = body(el, "MOBILE | VARCHAR(15) | NOT NULL | Mobile number")
el = body(el, "TOTALPRICE | Integer | NOT NULL | Amount paid (INR)")

# customercontactdb
el = body(el, "Table 3.7 - customercontactdb (Contact Enquiry)")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "CONTACTNAME | VARCHAR(100) | NOT NULL | Sender name")
el = body(el, "CONTACTNUMBER | VARCHAR(15) | NULLABLE | Phone number")
el = body(el, "CONTACTEMAIL | VARCHAR(100) | NOT NULL | Sender email")
el = body(el, "CONTACTSUBJECT | VARCHAR(200) | NOT NULL | Subject")
el = body(el, "CONTACTMESSAGE | TEXT | NOT NULL | Message content")

# ChatbotResponse
el = body(el, "Table 3.8 - ChatbotResponse")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "keywords | TEXT | NOT NULL | Keywords for pattern matching")
el = body(el, "response | TEXT | NOT NULL | Bot response text")
el = body(el, "query_type | VARCHAR(50) | NOT NULL | Query category")

# ChatbotConversation
el = body(el, "Table 3.9 - ChatbotConversation")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "username | VARCHAR(100) | NOT NULL | Participant username")
el = body(el, "user_message | TEXT | NOT NULL | User message")
el = body(el, "bot_response | TEXT | NOT NULL | Bot response")
el = body(el, "query_type | VARCHAR(50) | NOT NULL | Query category")
el = body(el, "created_at | DateTime | NOT NULL | Timestamp")

# Notification
el = body(el, "Table 3.10 - Notification")
el = body(el, "id | Integer | PRIMARY KEY, AUTO_INCREMENT | Unique identifier")
el = body(el, "username | VARCHAR(100) | NOT NULL | Target user")
el = body(el, "message | TEXT | NOT NULL | Notification content")
el = body(el, "is_read | Boolean | DEFAULT FALSE | Read status")
el = body(el, "created_at | DateTime | NOT NULL | Timestamp")

print("Data Dictionary (3.5) done")

# ============================================================
# STEP 3: Insert 3.6 ER Diagram after 3.5's last para
# ============================================================
# Find 3.6 heading to get its position
# 3.6 hasn't been inserted yet, so find 3.5 heading and its end
idx_35, p_35 = find_start('3.5 Data Dictionary')
s35_start, s35_end = section_range(idx_35)
anchor_35 = doc.paragraphs[s35_end - 1]
print(f"Last element of 3.5: P{s35_end-1}: {anchor_35.text.strip()[:60]}")

el = h2(anchor_35, "3.6 Entity Relationship Diagram")
el = body(el, "The Entity Relationship (ER) diagram illustrates the logical structure of the database with ten entities normalised up to Third Normal Form (3NF). The diagram below shows all entities, their attributes, primary and foreign keys, and the relationships between tables.")
el = body(el, "Figure 3.3 - Entity Relationship Diagram")
print("ER Diagram (3.6) done")

# ============================================================
# STEP 4: Insert 3.7 DFD after 3.6's last para
# ============================================================
idx_36, p_36 = find_start('3.6 Entity Relationship Diagram')
s36_start, s36_end = section_range(idx_36)
anchor_36 = doc.paragraphs[s36_end - 1]
print(f"Last element of 3.6: P{s36_end-1}: {anchor_36.text.strip()[:60]}")

el = h2(anchor_36, "3.7 Data Flow Diagram")
el = body(el, "A Data Flow Diagram (DFD) represents the flow of data through the Hotel Room Booking System at different levels of abstraction.")
el = body(el, "Context Diagram (Level 0 DFD): The context diagram shows the entire system as a single process with two external entities: Customer and Administrator. Data flows from Customer to System include: Registration Details, Login Credentials, Room Search Criteria, Booking Request, Payment Details, Contact Enquiry, and Chatbot Queries. Data flows from System to Customer include: Room Listings, Booking Confirmation, Payment Status, Search Results, Chatbot Responses, Email Notifications, and PDF Receipt. Data flows from Administrator to System include: Admin Login, Room Type Data, Room Data, Staff Data. Data flows from System to Administrator include: Booking Records, Payment Records, Contact Messages, and Dashboard Summary.")
el = body(el, "Level 1 DFD: The system is decomposed into eight processes: P1-User Management, P2-Room Management, P3-Search and Filtering, P4-Booking Engine, P5-Payment Processing (Razorpay integration), P6-Notification System (Email, OTP), P7-Chatbot Engine, P8-Admin Dashboard. Data stores include all ten database tables (D1-D10).")
el = body(el, "Figure 3.4 - Data Flow Diagram (Context and Level 1)")
print("DFD (3.7) done")

# ============================================================
# STEP 5: Insert Implementation of Security chapter (after Ch 7, before Ch 8)
# ============================================================
# Find Ch 8 heading
ch8_hdg = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    s = p.style.name
    if t.startswith('8. Implementation Plan') and 'Heading' in s:
        ch8_hdg = i
        break

# Find last para before Ch 8
anchor_ch8 = ch8_hdg - 1
while anchor_ch8 >= 0 and not doc.paragraphs[anchor_ch8].text.strip():
    anchor_ch8 -= 1
print(f"Ch 8 heading at P{ch8_hdg}, anchor P{anchor_ch8}: {doc.paragraphs[anchor_ch8].text.strip()[:60]}")

el = h1(doc.paragraphs[anchor_ch8], "8. Implementation of Security for the Software Developed")
el = body(el, "Security is a critical aspect of any web application handling customer data and financial transactions. The system implements multiple layers of security at application, database, and network levels.")
el = body(el, "8.1 User Authentication and Access Control: The system implements role-based access control with two user roles. Customer authentication uses server-side sessions with encrypted session data. Administrator access uses Django's built-in auth with PBKDF2 password hashing.")
el = body(el, "8.2 Password Security and OTP Verification: Password reset uses 5-digit OTP sent via email, stored server-side, single-use, with whitespace-stripped comparison.")
el = body(el, "8.3 CSRF Protection: All forms include Django CSRF tokens validated on every POST/PUT/DELETE request.")
el = body(el, "8.4 Payment Security: Razorpay (PCI DSS Level 1) handles all payment collection on secure checkout pages. Payment verification uses server-side signature validation.")
el = body(el, "8.5 SQL Injection Prevention: Django ORM parameterises all queries, preventing SQL injection.")
el = body(el, "8.6 XSS Prevention: Django template engine auto-escapes all variable output.")
el = body(el, "8.7 Data Integrity and Validation: Three-level validation (client, server, model). Foreign key constraints enforce referential integrity.")
el = body(el, "8.8 Secure Email Communication: SMTP with TLS encryption for booking confirmations and OTP delivery.")
print("Security chapter done")

# ============================================================
# STEP 6: Renumber chapters
# ============================================================
renames = [
    ("8. Implementation Plan", "9. Implementation Plan"),
    ("9. Limitations of the Project", "10. Limitations of the Project"),
    ("10. Future Application of the Project", "11. Future Application of the Project"),
    ("11. Conclusion", "12. Conclusion"),
    ("12. Bibliography", "13. Bibliography"),
]
for p in doc.paragraphs:
    t = p.text.strip()
    for old, new in renames:
        if t == old or t.startswith(old):
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
print("Chapter numbers fixed")

# Save
doc.save(SRC)
print("\nSaved! Verifying...")

# Verify
from docx import Document as D2
d2 = D2(SRC)
for i, p in enumerate(d2.paragraphs):
    t = p.text.strip()
    s = p.style.name
    if 'Heading' in s and t and (t[0].isdigit() or t.startswith('3.')):
        print(f"  P{i}: {t[:90]}")
