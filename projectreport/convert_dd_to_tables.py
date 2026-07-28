"""Replace plain-text Data Dictionary with proper Word tables + rebuild."""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import os

SRC = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx'

doc = Document(SRC)

# ============================================================
# Find and remove existing Data Dictionary paragraphs
# ============================================================
dd_start = None
dd_end = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == '3.5 Data Dictionary' and 'Heading' in p.style.name:
        dd_start = i
    if dd_start is not None and i > dd_start:
        if ('3.6' in t and 'Heading' in p.style.name) or ('4.' in t and 'Heading' in p.style.name):
            dd_end = i
            break

print(f"Data Dictionary: P{dd_start} to P{dd_end-1}")

# Collect the data from existing paragraphs before removing
dd_data = {}  # table_name -> [rows]
current_table = None
header_para = None  # the 3.5 heading para (keep this)

for i in range(dd_start, dd_end):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if i == dd_start:
        header_para = p
        continue
    if i == dd_start + 1:
        # intro text - keep it
        intro_para = p
        continue
    if t.startswith('Table 3.'):
        current_table = t
        dd_data[current_table] = []
    elif current_table and '|' in t:
        parts = [x.strip() for x in t.split('|')]
        dd_data[current_table].append(parts)

print(f"Found {len(dd_data)} tables with {sum(len(v) for v in dd_data.values())} total rows")

# Collect elements to remove (all from dd_start+2 to dd_end-1)
elements_to_remove = []
for i in range(dd_start + 2, dd_end):
    elements_to_remove.append(doc.paragraphs[i]._element)

# Remove them
BODY = doc.element.body
for e in elements_to_remove:
    try:
        BODY.remove(e)
    except:
        pass  # might already be detached

print(f"Removed {len(elements_to_remove)} plain-text paragraph elements")

# ============================================================
# Now insert proper tables after intro_para
# ============================================================
def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '333333')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def format_cell(cell, text, bold=False, font_size=9, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    run.bold = bold
    if color:
        run.font.color.rgb = color

# Insert after intro_para
insert_after = intro_para._element

headers = ['Field Name', 'Data Type', 'Constraints', 'Description']
col_widths = [Cm(3.5), Cm(3), Cm(4), Cm(6.5)]

for table_name, rows in dd_data.items():
    # Insert table name as heading paragraph
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = table_name
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    insert_after.addnext(new_p)
    insert_after = new_p

    # Create the table
    num_rows = len(rows) + 1  # +1 for header
    table = doc.add_table(rows=num_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width
    
    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        format_cell(cell, h, bold=True, font_size=9)
        set_cell_shading(cell, '2F5496')
        # White text for header
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = None  # Reset
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for ri, row_data in enumerate(rows):
        for cj, cell_text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[cj]
            format_cell(cell, cell_text, font_size=9)
            # Alternating row colors
            if ri % 2 == 0:
                set_cell_shading(cell, 'D6E4F0')
    
    # Move table XML after the heading paragraph
    table_elem = table._element
    insert_after.addnext(table_elem)
    insert_after = table_elem

print("Tables inserted!")

# Save
doc.save(SRC)
print("Saved!")
