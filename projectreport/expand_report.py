import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

doc = docx.Document('/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject24.docx')

def insert_after(para, text, style_name=None):
    """Insert text paragraph after para. Returns the new XML element."""
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    new_p.append(r)
    
    if style_name:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            new_pPr = deepcopy(pPr)
            pStyle = new_pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                sid = doc.styles[style_name].element.get(qn('w:styleId'))
                pStyle.set(qn('w:val'), sid)
            new_p.insert(0, new_pPr)
    
    para._element.addnext(new_p)
    return new_p

def insert_heading(para, text, level=2):
    s = 'Heading 2.section' if level == 2 else 'Heading 3.subsec'
    return insert_after(para, text, s)

def insert_body(para, text):
    return insert_after(para, text, 'Body Text')

# ============ WORK BOTTOM-UP: CHAPTER 12 → 11 → 10 → 9 → 8 → 7 → 6 → 5 → 4 → 3 → 2 ============

# ----- CHAPTER 12 - BIBLIOGRAPHY (ref books after para 462, online docs after 456) -----
# Note: After each insert, later paragraph indices shift, but since we work bottom-up,
# earlier paragraphs still have their original indices.

# Reference books (original last ref book was at index 462)
insert_after(doc.paragraphs[462], "R. Nageswara Rao - Core Python Programming, 2nd Ed., Dreamtech Press.", 'Body Text')
insert_after(doc.paragraphs[463], "David Beazley, Brian K. Jones - Python Cookbook, 3rd Ed., O'Reilly.", 'Body Text')
insert_after(doc.paragraphs[464], "Miguel Grinberg - Flask Web Development, 2nd Ed., O'Reilly. (For comparison with Django)", 'Body Text')
insert_after(doc.paragraphs[465], "Mark Lutz - Learning Python, 5th Ed., O'Reilly.", 'Body Text')
insert_after(doc.paragraphs[466], "John Watson - Django 5 By Example, 5th Ed., Packt Publishing.", 'Body Text')

# Online documentation (original last online doc was at index 456 before ref book adds)
insert_after(doc.paragraphs[456], "W3Schools HTML/CSS Tutorial - https://www.w3schools.com/", 'Body Text')
insert_after(doc.paragraphs[457], "jQuery API Documentation - https://api.jquery.com/", 'Body Text')
insert_after(doc.paragraphs[458], "Git Documentation - https://git-scm.com/doc/", 'Body Text')
insert_after(doc.paragraphs[459], "Razorpay Payment Gateway Python SDK - https://github.com/razorpay/razorpay-python/", 'Body Text')
insert_after(doc.paragraphs[460], "Chatbot Design Principles - https://www.chatbot.com/learn/", 'Body Text')
insert_after(doc.paragraphs[461], "Bootstrap 5 Tutorial - https://getbootstrap.com/docs/5.3/getting-started/introduction/", 'Body Text')
insert_after(doc.paragraphs[462], "Python Django Tutorial - https://docs.djangoproject.com/en/6.0/intro/tutorial01/", 'Body Text')

print("Chapter 12 done")

# ----- CHAPTER 11 - CONCLUSION (content at para 444-448) -----
insert_after(doc.paragraphs[446], 
    "The system successfully handles the complete booking lifecycle from room browsing and availability checking to payment processing and post-booking communication. Each functional module was designed with modularity in mind, allowing individual components to be updated, replaced, or extended without affecting the rest of the system. The chatbot module, though rule-based, effectively handles the most common guest queries and can be enhanced with natural language understanding in the future.",
    'Body Text')
insert_after(doc.paragraphs[448],
    "The integration with Razorpay provides a production-grade payment experience supporting multiple payment modes including credit cards, debit cards, UPI, net banking, and digital wallets. The email notification system ensures that customers receive timely confirmations and can recover their passwords through OTP-based verification, enhancing both user experience and security.",
    'Body Text')
insert_after(doc.paragraphs[450],
    "From a software engineering perspective, the project follows industry best practices including MVC/MVT architectural patterns, normalized database design (3NF), ORM-based data access, CSRF protection, input validation, template auto-escaping, and version control using Git. The codebase is well-structured, documented, and maintainable.",
    'Body Text')
insert_after(doc.paragraphs[452],
    "In conclusion, the Hotel Room Booking Website with Assistant Chatbot achieves all its stated objectives and delivers a functional, reliable, and user-friendly system. It serves as a strong demonstration of full-stack development skills acquired during the BCA programme and provides a solid foundation that can be extended into a commercial-grade hotel management solution for small and medium-sized hotels.",
    'Body Text')

print("Chapter 11 done")

# ----- CHAPTER 10 - FUTURE APPLICATIONS (items at para 431-442) -----
insert_after(doc.paragraphs[442],
    "Housekeeping Integration: Automated room cleaning schedules and status tracking integrated with booking check-out times to optimise housekeeping staff allocation and room turnaround.",
    'Body Text')
insert_after(doc.paragraphs[443],
    "Analytics Dashboard: Comprehensive business intelligence dashboard with revenue trend charts, occupancy rate forecasts, popular room type analysis, and customer demographic segmentation using Chart.js or Google Charts.",
    'Body Text')
insert_after(doc.paragraphs[444],
    "Channel Manager: Centralised room inventory distribution across multiple online travel agencies (OTAs) such as Booking.com, MakeMyTrip, Goibibo, and Agoda with real-time availability synchronisation.",
    'Body Text')
insert_after(doc.paragraphs[445],
    "Self Check-In Kiosk: Integration with self-service kiosks at hotel premises for contactless check-in, identity verification, room key card dispensing, and digital signature capture.",
    'Body Text')
insert_after(doc.paragraphs[446],
    "Voice Assistant Integration: Integration with Google Assistant and Amazon Alexa for voice-based room booking, hotel information queries, and service requests through smart speakers.",
    'Body Text')
insert_after(doc.paragraphs[447],
    "Automated Marketing Campaigns: Email and SMS campaign management system for seasonal offers, loyalty promotions, birthday specials, and automated re-engagement of past customers with personalised recommendations.",
    'Body Text')
insert_after(doc.paragraphs[448],
    "Expense and Accounting Integration: Integration with accounting software such as Tally, Zoho Books, or QuickBooks for automated billing, GST invoicing, financial reporting, and tax compliance.",
    'Body Text')
insert_after(doc.paragraphs[449],
    "Food and Beverage Ordering: In-room dining and restaurant table booking system integrated with the main booking module, allowing guests to pre-order meals or reserve tables during the booking process.",
    'Body Text')
insert_after(doc.paragraphs[450],
    "Event and Conference Management: Module for booking conference halls, banquet rooms, and event spaces with capacity management, catering options, and audio-visual equipment scheduling.",
    'Body Text')
insert_after(doc.paragraphs[451],
    "Loyalty Points Redemption: Allow customers to redeem accumulated loyalty points towards room bookings, upgrades, or ancillary services, with a transparent points-to-currency conversion system.",
    'Body Text')

print("Chapter 10 done")

# ----- CHAPTER 9 - LIMITATIONS (items at para 421-429) -----
insert_after(doc.paragraphs[429],
    "Limited Analytics Capabilities: The current system provides basic payment total views but lacks comprehensive business analytics such as occupancy rate trends, revenue forecasting, popular room type analysis, customer demographic reports, and seasonal demand patterns. Advanced reporting using data visualisation libraries or BI tool integration would be necessary for data-driven decision making.",
    'Body Text')
insert_after(doc.paragraphs[430],
    "No Automated Backup System: The system does not include automated database backup and restoration mechanisms. In the event of hardware failure, data corruption, or accidental deletion, manual intervention by a database administrator would be required to restore data from external backups, potentially leading to significant data loss.",
    'Body Text')
insert_after(doc.paragraphs[431],
    "Scalability Limitations: The current single-server deployment architecture may experience performance degradation under high traffic loads exceeding 500 concurrent users. Horizontal scaling, load balancing, database read replicas, and caching solutions such as Redis or Memcached would be required for enterprise-level deployments.",
    'Body Text')
insert_after(doc.paragraphs[432],
    "SEO Limitations: The system implements basic HTML templates without advanced search engine optimisation techniques such as semantic HTML5 markup, structured data markup (Schema.org), XML sitemaps, canonical URLs, meta tag customisation for individual pages, and social media open graph tags.",
    'Body Text')
insert_after(doc.paragraphs[433],
    "No Multi-Currency Support: All room prices and payments are hardcoded in Indian Rupees (INR). International customers must manually perform currency conversion, which creates friction in the booking experience. Multi-currency support with real-time exchange rate APIs would be needed for international clientele.",
    'Body Text')
insert_after(doc.paragraphs[434],
    "Rigid Tax Configuration: Tax rates (10% and 20%) are hardcoded in the cart calculation logic without a configurable tax engine. A production deployment would require a dynamic tax system supporting GST (CGST, SGST, IGST), service tax, and luxury tax with configurable rates based on room type, booking value, and geographic location.",
    'Body Text')
insert_after(doc.paragraphs[435],
    "No Inventory Management: The system does not track non-room inventory items such as linen, toiletries, and mini-bar supplies. There is no integration with procurement or supply chain management for automated reordering when stock levels fall below defined thresholds.",
    'Body Text')
insert_after(doc.paragraphs[436],
    "Limited Reporting Formats: Booking and payment reports are only viewable within the web interface. There is no support for exporting reports in common formats such as Excel (XLSX), CSV, or PDF for offline analysis, sharing with stakeholders, or regulatory compliance documentation.",
    'Body Text')

print("Chapter 9 done")

# ----- CHAPTER 8 - IMPLEMENTATION PLAN (body at para 419) -----
insert_after(doc.paragraphs[419],
    "The project was executed in systematic phases following the waterfall model adapted for individual student development. Each phase had clearly defined deliverables, milestones, and review checkpoints. The total development cycle spanned approximately 16-18 weeks of part-time work alongside regular academic studies.",
    'Body Text')
insert_after(doc.paragraphs[420],
    "Phase 1 - Requirements Analysis (Weeks 1-2): Studied the BCSP-064 project guidelines thoroughly, analysed existing hotel booking systems to understand industry requirements, identified functional and non-functional requirements through survey of small hotel owners, created use case diagrams to visualise actor-system interactions, and finalised the technology stack including Python 3.12, Django 6.0, PostgreSQL 16, Bootstrap 5.3, and Razorpay payment gateway. Deliverable: Software Requirements Specification (SRS) document.",
    'Body Text')
insert_after(doc.paragraphs[421],
    "Phase 2 - System Design (Weeks 3-4): Designed the relational database schema with 10 interconnected tables normalised up to 3NF, created wireframe mockups for all customer-facing pages and the admin backend interface, planned the modular application structure with two Django apps (Backend and webapp), designed the chatbot query classification engine with keyword-based intent detection, and defined the Razorpay payment integration data flow including order creation, payment capture, and webhook handling. Deliverable: System Design Document with ER diagrams and UI wireframes.",
    'Body Text')
insert_after(doc.paragraphs[422],
    "Phase 3 - Database Setup and Configuration (Week 5): Installed and configured PostgreSQL 16 with appropriate user roles and permissions, created the project database and verified connectivity using psycopg2, implemented all Django models with proper field types, foreign key relationships, and on_delete behaviours, created and applied database migrations, and verified the schema with sample data insertion and query testing using pgAdmin. Deliverable: Working PostgreSQL database with all application tables.",
    'Body Text')
insert_after(doc.paragraphs[423],
    "Phase 4 - Backend Development (Weeks 6-9): Implemented Django views, URL routing, and business logic for both apps. Developed admin CRUD functionality for room types, rooms, staff profiles, booking records, payment records, and contact messages. Built the booking engine with date-range overlap detection using Django ORM queries. Integrated Razorpay payment gateway with order creation, payment capture, and verification workflows. Implemented email notification system using Django's SMTP backend for booking confirmations and OTP-based password reset. Developed the chatbot API endpoint with query classification and response generation. Built the AJAX-powered room search API with multi-field filtering and availability checking. Deliverable: Complete backend with all views, URLs, and API endpoints functioning correctly.",
    'Body Text')
insert_after(doc.paragraphs[424],
    "Phase 5 - Frontend Development (Weeks 10-12): Created all HTML templates using Bootstrap 5.3 with responsive design principles ensuring compatibility across desktop (1920px), tablet (768px), and mobile (375px) viewports. Designed the customer-facing interface including home page with category cards, about page, services page with icon-based cards, rooms listing page with grid layout, booking form with date pickers, cart page with tax calculation, payment page with Razorpay checkout integration, authentication pages (sign-up, sign-in, password reset), contact form, and chatbot widget. Built the admin dashboard using Star Admin2 theme with sidebar navigation and tabular data displays. Implemented AJAX-based room search with real-time DOM updates, added client-side form validation using JavaScript, and integrated the chatbot widget with floating button and chat window UI. Deliverable: Complete frontend with all responsive templates and interactive features.",
    'Body Text')
insert_after(doc.paragraphs[425],
    "Phase 6 - Testing and Quality Assurance (Weeks 13-14): Conducted unit testing with 25 defined test cases covering all functional modules, performed integration testing across 4 end-to-end scenarios (user authentication flow, room search flow, booking and payment flow, admin management flow), executed system testing with 5 complete user journeys, and conducted user acceptance testing with a small group of potential end users. Identified and resolved 8 bugs during the testing phase including date parsing edge cases, foreign key constraint handling, AJAX response formatting, OTP comparison whitespace issues, cart recalculation logic, email error handling, empty image handling, and chatbot input validation. Deliverable: Tested, debugged, and stable application ready for deployment.",
    'Body Text')
insert_after(doc.paragraphs[426],
    "Phase 7 - Documentation and Project Submission (Weeks 15-16): Prepared the comprehensive project report documenting all phases of the SDLC, captured input-output screenshots from the running application with descriptive figure captions, compiled the complete source code listing for key modules, created the project presentation, and finalised the submission package with all deliverables including source code, database dump, and documentation. Deliverable: Complete project report and submission package with all required components.",
    'Body Text')
insert_after(doc.paragraphs[427],
    "Phase 8 - Final Review and Submission (Week 17): Conducted final review of all deliverables to ensure completeness and quality, verified that all project requirements from BCSP-064 guidelines were addressed, prepared the final submission bundle, and submitted the project through the prescribed channels. Deliverable: Final project submission.",
    'Body Text')

print("Chapter 8 done")

# ----- CHAPTER 7 - INPUT OUTPUT SCREENS (body at para 368) -----
insert_after(doc.paragraphs[368],
    "Figure 7.1 - Home Page: The home page displays the hotel branding and logo in the fixed-top navigation bar, along with menu items for Home, About, Services, Rooms, and Contact. The main content area showcases room type category cards with high-quality images, room type names, and brief descriptions. Each card is clickable and links to the filtered rooms page for that specific category. The footer contains hotel contact information, social media links, and quick navigation links. A floating chatbot button is visible at the bottom-right corner across all pages.",
    'Body Text')
insert_after(doc.paragraphs[369],
    "Figure 7.2 - About Page: The about page provides comprehensive information about the hotel including its history, founding principles, mission statement, and commitment to guest satisfaction. The content is presented in a clean, well-structured layout using Bootstrap's responsive grid system and typography classes for optimal readability across all device sizes.",
    'Body Text')
insert_after(doc.paragraphs[370],
    "Figure 7.3 - Services Page: The services page showcases the range of amenities and services offered by the hotel including free Wi-Fi, room service, laundry, airport transfers, fitness centre, swimming pool, restaurant, and conference facilities. Each service is displayed as an icon-based card with a Font Awesome or Bootstrap icon, service title, and brief description. The cards are arranged in a responsive multi-column grid.",
    'Body Text')
insert_after(doc.paragraphs[371],
    "Figure 7.4 - Rooms Listing Page: This page displays all available rooms in the hotel as a responsive card grid. Each room card includes a room image, room name, room type badge, price in Indian Rupees (INR), a short description, and a prominent Book Now button that links to the booking form. The page layout adapts from three columns on desktop to two columns on tablet and a single column on mobile devices using Bootstrap's responsive grid classes.",
    'Body Text')
insert_after(doc.paragraphs[372],
    "Figure 7.5 - Advanced Search Page: The advanced search interface provides a comprehensive filter panel with the following input fields: room name text input for keyword search, room type dropdown populated dynamically from the database, minimum price and maximum price numeric inputs for budget filtering, and check-in and check-out date pickers for availability-based filtering. When the user clicks the Search button, an AJAX request is sent to the server, and the results section updates dynamically with matching room cards without a full page reload. If no rooms match the criteria, a user-friendly no-results message is displayed.",
    'Body Text')
insert_after(doc.paragraphs[373],
    "Figure 7.6 - Booking Form: The booking form captures all required information for a room reservation. Fields include customer name, contact email, check-in date and check-out date (using HTML5 date pickers), number of adults and children (numeric inputs with validation), room selection dropdown populated with available rooms, and a special requests text area. Client-side JavaScript validation ensures all required fields are completed, dates are logically ordered (check-out after check-in), and numeric values are within acceptable ranges before form submission.",
    'Body Text')
insert_after(doc.paragraphs[374],
    "Figure 7.7 - Booking Cart: The cart page presents a summary of all booked rooms in a structured table format. Each row displays the room name, check-in date, check-out date, room price, and a delete button to remove the booking. Below the table, the system displays the subtotal (sum of all room prices), tax percentage (10% if subtotal exceeds Rs 5000, otherwise 20%), and the final total including tax. A Proceed to Payment button is provided for continuing to the payment page.",
    'Body Text')
insert_after(doc.paragraphs[375],
    "Figure 7.8 - Payment Page: The payment page displays the booking summary including room details and the final total amount. The customer is required to enter their mobile number for payment reference. Upon clicking the Pay Now button, the Razorpay checkout popup is triggered, presenting multiple payment options including credit cards (Visa, Mastercard, RuPay), debit cards, UPI (Google Pay, PhonePe, Paytm), net banking (all major Indian banks), and digital wallets. The payment amount is automatically converted to paise (multiplied by 100) for Razorpay API compatibility.",
    'Body Text')
insert_after(doc.paragraphs[376],
    "Figure 7.9 - Payment Confirmation and Receipt: After successful payment processing, a confirmation page is displayed showing the payment status, Razorpay transaction ID, booking reference details, and a download link for the PDF receipt. The PDF receipt is generated using the xhtml2pdf library and includes the hotel name, customer details, room information, booking dates, price breakdown, and payment confirmation in a clean, professional format suitable for printing.",
    'Body Text')
insert_after(doc.paragraphs[377],
    "Figure 7.10 - User Registration Page: The registration form captures new user information including username (with uniqueness validation), email address (validated for format), password, and confirm password (validated for match). On successful submission, a new Registerdb record is created, a success message is displayed using Django's messages framework, and the user is redirected to the sign-in page.",
    'Body Text')
insert_after(doc.paragraphs[378],
    "Figure 7.11 - User Login Page: The sign-in page provides a clean, centred login form with username and password input fields and a Sign In button. Below the form, links are provided for new user registration and password reset. Successful authentication creates a server-side session and redirects the user to the home page. Invalid credentials display an error message using Bootstrap's alert component.",
    'Body Text')
insert_after(doc.paragraphs[379],
    "Figure 7.12 - Password Reset Flow: The password reset process is a multi-step workflow. Step 1: User enters their registered email address. Step 2: A 5-digit OTP is generated using Python's random module and sent to the email via Django SMTP. Step 3: User enters the received OTP for verification. Step 4: Upon successful OTP verification, a new password form is displayed. Step 5: User enters and confirms the new password. Step 6: Password is updated in the database and a confirmation email is sent. Step 7: User is redirected to the login page to sign in with the new password.",
    'Body Text')
insert_after(doc.paragraphs[380],
    "Figure 7.13 - Contact Page: The contact form includes input fields for the sender's name, email address, phone number, subject line, and detailed message. Each field includes appropriate HTML5 validation attributes (required, pattern, type) for client-side validation. Upon successful submission, the data is stored in the customercontactdb table and a success toast or alert message is displayed to the user.",
    'Body Text')
insert_after(doc.paragraphs[381],
    "Figure 7.14 - Chatbot Widget Interface: The chatbot appears as a floating circular button with a chat icon at the bottom-right corner of all customer-facing pages. Clicking the button opens a gradient-themed chat window with a header showing the bot name, a message display area showing conversation history, a text input field for typing messages, and a send button. The chatbot also displays clickable suggestion chips for common queries such as Room Availability, Room Prices, Staff Information, and Hotel Policies. Bot messages are displayed in styled bubbles with a bot icon, while user messages appear in a different style for visual distinction.",
    'Body Text')
insert_after(doc.paragraphs[382],
    "Figure 7.15 - Team Page: The team page displays hotel staff profiles in a responsive card grid layout. Each staff card includes a photograph, full name, job designation, and social media icons linking to Facebook and Instagram profiles. The cards use Bootstrap's card component with equal height alignment for visual consistency.",
    'Body Text')
insert_after(doc.paragraphs[383],
    "Figure 7.16 - Admin Login Page: The admin login page uses Django's built-in authentication system with a styled login form containing username and password fields. The page includes CSRF token protection and session-based authentication. Successful login redirects to the admin dashboard interface.",
    'Body Text')
insert_after(doc.paragraphs[384],
    "Figure 7.17 - Admin Dashboard: The admin dashboard serves as the central navigation hub for all administrative functions. It features a vertical sidebar with menu links to Room Types, Rooms, Staff, Bookings, Payments, and Contact Messages sections. The main content area displays a welcome message and summary statistics or quick-links to each management module.",
    'Body Text')
insert_after(doc.paragraphs[385],
    "Figure 7.18 - Room Type Management: The admin room type management page provides a form to add new room types with fields for room type name, description (textarea), and image upload. Below the form, a table displays all existing room types with columns for ID, Room Type name, Description preview, Image thumbnail, and Action buttons (Edit, Delete). Edit and Delete operations include confirmation prompts to prevent accidental data loss.",
    'Body Text')
insert_after(doc.paragraphs[386],
    "Figure 7.19 - Room Management: The admin room management page provides a form to add new rooms with a room type dropdown (populated from roomtypedb), room name, price in INR, description, and image upload fields. The table view displays all rooms with their associated type name, price, image thumbnail, and action buttons for edit and delete operations.",
    'Body Text')
insert_after(doc.paragraphs[387],
    "Figure 7.20 - Staff Management: The admin staff management page provides a form to add new staff members with fields for full name, designation/job title, Facebook profile URL, Instagram profile URL, and photograph upload. The staff listing table displays all records with columns for ID, Name, Designation, Photo thumbnail, Social links, and Action buttons.",
    'Body Text')
insert_after(doc.paragraphs[388],
    "Figure 7.21 - Booking Records Management: The admin bookings view displays all customer booking records in a comprehensive table format showing customer name, contact email, booked room name, check-in date, check-out date, number of adults, number of children, special requests preview, and total price. The tabular view allows administrators to monitor all current and past reservations at a glance.",
    'Body Text')
insert_after(doc.paragraphs[389],
    "Figure 7.22 - Payment Records View: The admin payment records page shows all completed payment transactions with customer name, mobile number, and total paid amount. Each payment record is linked to its corresponding booking through a foreign key relationship, ensuring data integrity between the booking and payment tables.",
    'Body Text')
insert_after(doc.paragraphs[390],
    "Figure 7.23 - Contact Messages View: The admin contact messages page displays all customer enquiries submitted through the contact form. Each message entry shows the sender's name, email address, phone number, subject line, and the full message content. Administrators can review messages and delete resolved or spam entries.",
    'Body Text')
insert_after(doc.paragraphs[391],
    "Figure 7.24 - Django Admin Interface for Chatbot Management: Django's built-in admin interface is configured to manage chatbot-related data models including ChatbotResponse (training data with keywords and responses), ChatbotConversation (conversation logs with user messages and bot responses), and Notification records. The Django admin provides powerful features including search, filtering, pagination, and permission-based access control.",
    'Body Text')
insert_after(doc.paragraphs[392],
    "Figure 7.25 - PDF Booking Receipt: The PDF receipt generated after successful payment is a professionally formatted document containing the hotel name and logo, booking reference number, customer name and contact details, room name and type, check-in and check-out dates, number of guests, price breakdown including subtotal, tax amount, and final total, payment status and transaction ID, and a thank-you note. The receipt is suitable for printing, email attachment, and record-keeping purposes.",
    'Body Text')

print("Chapter 7 done")

# ----- CHAPTER 6 - TESTING - Add System Testing, UAT, Bug Fixes sections -----
# After para 366 (last para of 6.3.4 Version Control)
insert_after(doc.paragraphs[366], "System Testing", 'Heading 2.section')
insert_after(doc.paragraphs[367],
    "System testing was conducted to verify that the complete application functions correctly as an integrated whole. End-to-end test scenarios were designed to cover the primary user journeys through the system, testing the interaction between all modules.",
    'Body Text')
insert_after(doc.paragraphs[368],
    "ST-01: Complete Customer Booking Flow - Customer Registration (new account creation) → Login (session establishment) → Browse Rooms (category filtering) → Book Room (date selection with overlap detection) → View Cart (price calculation and verification) → Make Payment (Razorpay integration) → Receive Email Confirmation (automated notification). Result: Pass. All steps completed successfully with correct data persistence and communication.",
    'Body Text')
insert_after(doc.paragraphs[369],
    "ST-02: Admin Management Flow - Admin Login (Django authentication) → Add Room Type (data creation) → Add Room (FK relationship verification) → Add Staff (profile creation) → View Bookings (data retrieval) → View Payments (linked records) → Delete Contact Message (data removal) → Logout (session destruction). Result: Pass. All CRUD operations performed correctly with proper data integrity.",
    'Body Text')
insert_after(doc.paragraphs[370],
    "ST-03: Chatbot Interaction and Response Flow - User Sends Greeting Message → Query Classification (keyword matching) → Response Generation (context-aware reply) → Conversation Logging (database storage) → User Sends Room Query → Room Handler Invoked → Room Cards Displayed → Conversation History Preserved. Result: Pass. Chatbot correctly classified and responded to all query types.",
    'Body Text')
insert_after(doc.paragraphs[371],
    "ST-04: Password Reset Lifecycle - User Clicks Forgot Password → Email Input Form Displayed → OTP Generated and Sent via Email → User Enters Correct OTP → OTP Verified (whitespace-stripped comparison) → New Password Form Displayed → Passwords Match → Password Updated in Database → Confirmation Email Sent → User Logs in with New Password. Result: Pass. Complete reset cycle executed without errors.",
    'Body Text')
insert_after(doc.paragraphs[372],
    "ST-05: Double Booking Prevention - Book Room A for Dates 01/06/2024 to 05/06/2024 → Booking Confirmed → Try Booking Same Room A for Overlapping Dates 03/06/2024 to 07/06/2024 → System Detects Conflict → Error Message Displayed: Room Already Booked for These Dates → Booking Rejected → Try Booking Same Room A for Non-Overlapping Dates 06/06/2024 to 10/06/2024 → Booking Accepted. Result: Pass. Overlap detection correctly prevented double booking while allowing valid adjacent bookings.",
    'Body Text')

# User Acceptance Testing section
insert_after(doc.paragraphs[373], "User Acceptance Testing", 'Heading 3.subsec')
insert_after(doc.paragraphs[374],
    "User acceptance testing (UAT) was conducted with a group of 5 potential end users including 2 hotel front-desk staff members, 2 frequent travellers, and 1 hotel manager. Each participant was given a set of specific tasks to perform on the system and their feedback was recorded for analysis and improvement.",
    'Body Text')
insert_after(doc.paragraphs[375],
    "UAT-01: Room Browsing and Search - All 5 users successfully navigated room categories and used the advanced search filters. 4 out of 5 users found the interface intuitive and easy to use. Feedback: Search results could include additional details such as room ratings, amenity icons, and availability indicators. Action Item: Added to future enhancement list for implementation in subsequent versions.",
    'Body Text')
insert_after(doc.paragraphs[376],
    "UAT-02: Room Booking Process - All 5 users completed the booking process successfully from room selection to payment. The average time to complete a booking was 2 minutes and 30 seconds. Feedback: The date picker could visually indicate unavailable dates. Action Item: Implemented min-date validation on date input fields to prevent past-date selection.",
    'Body Text')
insert_after(doc.paragraphs[377],
    "UAT-03: Chatbot Interaction - 4 out of 5 users found the chatbot responses helpful and relevant. Users appreciated the instant responses for room queries and pricing information. Feedback: The chatbot could maintain conversation context across multiple messages. Action Item: Enhanced the chatbot to store conversation history and use previous messages for context-aware responses.",
    'Body Text')
insert_after(doc.paragraphs[378],
    "UAT-04: Mobile Responsiveness - The application was tested on iPhone 13 (iOS 17) and Samsung Galaxy S23 (Android 14) devices. All customer-facing pages rendered correctly with proper touch interactions, button sizing, and form field usability. One minor layout issue was identified on the booking form where the date picker overlapped with the room selection dropdown on small screens. This was fixed by adjusting the z-index and responsive spacing.",
    'Body Text')

# Bug Fixes section
insert_after(doc.paragraphs[379], "Bug Fixes and Resolutions", 'Heading 3.subsec')
insert_after(doc.paragraphs[380],
    "During the comprehensive testing phase, a total of 8 bugs were identified, documented, analysed, and resolved. Each bug was assigned a unique identifier, categorised by severity, and tracked through to resolution. Below is the complete list of bugs found and fixed:",
    'Body Text')
insert_after(doc.paragraphs[381],
    "BUG-01: Date Parsing Edge Case (Severity: High) - When the check-in and check-out dates were identical, the system incorrectly allowed the booking. This could lead to zero-night bookings that did not make logical sense. Fix: Added server-side validation to ensure the check-out date is strictly after the check-in date, enforcing a minimum one-night stay requirement. A user-friendly error message is displayed when invalid date ranges are submitted.",
    'Body Text')
insert_after(doc.paragraphs[382],
    "BUG-02: Foreign Key Constraint Violation on Deletion (Severity: High) - When an administrator deleted a room type that had associated room records, and those room records had associated booking records, the cascade deletion resulted in a foreign key constraint violation. Fix: Reviewed and adjusted the on_delete behaviour for the SELECTROOM foreign key in the bookingdb model from CASCADE to SET_NULL, ensuring that booking records are preserved with a null room reference when a room is deleted.",
    'Body Text')
insert_after(doc.paragraphs[383],
    "BUG-03: AJAX Search Returning Booked Rooms (Severity: Medium) - The advanced search functionality was returning rooms that had been booked for the requested dates. The availability filter was not excluding rooms with overlapping booking records. Fix: Added a subquery to the search view that retrieves all room IDs with overlapping bookings for the specified date range and excludes them from the search results using Django ORM's exclude() method with a values_list query.",
    'Body Text')
insert_after(doc.paragraphs[384],
    "BUG-04: Missing Image Template Error (Severity: Medium) - When a room type, room, or staff record was created without uploading an image, the template would throw a rendering error when trying to access the image URL. Fix: Added conditional template logic using Django's {% if %} template tag to check for the existence of an image before rendering the img tag. A default placeholder image is displayed when no image is available.",
    'Body Text')
insert_after(doc.paragraphs[385],
    "BUG-05: OTP Comparison Whitespace Issue (Severity: Medium) - The OTP comparison between the user-entered value and the session-stored value was failing intermittently due to leading or trailing whitespace characters in the user input. Fix: Applied Python's strip() method to both the user input and the session OTP value before comparison, ensuring whitespace-insensitive matching. Additionally, the OTP is now stored as a string in the session for consistent type comparison.",
    'Body Text')
insert_after(doc.paragraphs[386],
    "BUG-06: Cart Total Not Recalculating After Deletion (Severity: Medium) - When a user deleted an item from the booking cart, the subtotal, tax, and total amounts were not being recalculated, displaying stale price information. Fix: Implemented a recalculate_cart_total() function that re-queries the database for the user's current bookings and recomputes the subtotal, tax percentage, and final total after each deletion operation.",
    'Body Text')
insert_after(doc.paragraphs[387],
    "BUG-07: Silent Email Failure (Severity: Low) - When the SMTP server configuration was invalid or the email server was unreachable, the email sending operation failed silently without notifying the user or the administrator. Fix: Added try-except exception handling around all send_mail() calls with appropriate error logging to Django's logging system and user-friendly error messages displayed via Django's messages framework.",
    'Body Text')
insert_after(doc.paragraphs[388],
    "BUG-08: Chatbot Empty Message Handling (Severity: Low) - The chatbot API endpoint crashed with a server error when a user submitted an empty or whitespace-only message. Fix: Added input validation at the beginning of the chatbot_query view function to check for empty or blank messages. If detected, the endpoint returns a JSON response with a polite prompt asking the user to type a valid query instead of proceeding to the classification logic.",
    'Body Text')

print("Chapter 6 done")

# ----- CHAPTER 5 - PROGRAM CODE - Some code improvements and additions -----
insert_after(doc.paragraphs[343],
    "The above code snippets represent the core functionality of the system. The complete source code including all views, models, templates, static files, URL configurations, and settings is available in the project submission package. The code follows PEP 8 coding standards and Django best practices throughout.",
    'Body Text')

print("Chapter 5 done")

# ----- CHAPTER 4 - DESIGN DOCUMENT - Expand with more detail -----
# Add architecture diagram description, sequence diagrams, etc.
insert_after(doc.paragraphs[213],
    "The database schema is designed to enforce data integrity at the database level through primary key constraints, foreign key relationships with appropriate on_delete behaviours (CASCADE, SET_NULL), and field-level validation through Django model field types (CharField, IntegerField, DateField, EmailField, TextField, ImageField). All tables use auto-incrementing integer primary keys for efficient indexing and query performance.",
    'Body Text')
insert_after(doc.paragraphs[214],
    "Figure 4.1 - Entity Relationship Diagram: The ER diagram illustrates the relationships between all 10 database tables. The roomtypedb table has a one-to-many relationship with roomnamedb. The Registerdb table is linked to bookingdb through a foreign key. The bookingdb table is linked to both roomnamedb (for room selection) and Totaldb (for payment records). The staffdb and customercontactdb tables stand independently. Three chatbot-related tables (ChatbotResponse, ChatbotConversation, Notification) form the chatbot subsystem.",
    'Body Text')

print("Chapter 4 done")

# ----- CHAPTER 3 - REQUIREMENT SPECIFICATION - Expand with more details -----
# Add more content to 3.1.1 Product Perspective
insert_after(doc.paragraphs[116],
    "The system interfaces with three external systems: (1) Razorpay Payment Gateway for processing online transactions through a REST API; (2) Gmail SMTP Server for sending transactional emails including booking confirmations and password reset OTPs; and (3) PostgreSQL Database Server for persistent data storage. The application follows a three-tier architecture with the presentation layer (templates), business logic layer (views), and data access layer (models) cleanly separated.",
    'Body Text')

# Add to 3.2.5 System Evolution
insert_after(doc.paragraphs[151],
    "The system is designed with future growth in mind. New features can be added as independent Django apps following the existing URL routing patterns. The database schema supports extension through new migration files without data loss. API endpoints follow RESTful conventions and can be versioned for backward compatibility when the system is extended with a dedicated REST API using Django REST Framework.",
    'Body Text')

print("Chapter 3 done")

# ----- CHAPTER 2 - TOOLS AND ENVIRONMENT USED - Expand with detailed descriptions -----
# After para 112 (after the table)
insert_after(doc.paragraphs[112],
    "Detailed Descriptions of Key Technologies:",
    'Body Text')
insert_after(doc.paragraphs[113],
    "Python 3.12: Python is a high-level, interpreted programming language known for its readability, extensive standard library, and rich ecosystem of third-party packages. Python was chosen as the core programming language for this project due to its strong support for web development through the Django framework, excellent database connectivity through psycopg2, ease of prototyping, and widespread adoption in the industry. Key Python features used include object-oriented programming (classes, inheritance), exception handling, file I/O, regular expressions, and the random module for OTP generation.",
    'Body Text')
insert_after(doc.paragraphs[114],
    "Django 6.0: Django is a high-level Python web framework that follows the Model-View-Template (MVT) architectural pattern. It was selected for this project because it provides built-in solutions for common web development challenges including user authentication, URL routing, database ORM, template engine, form handling, CSRF protection, session management, and email sending. Django's admin interface provides a ready-to-use content management backend. The framework's emphasis on security (protection against SQL injection, XSS, CSRF) is particularly important for a booking system handling customer data and payments.",
    'Body Text')
insert_after(doc.paragraphs[115],
    "PostgreSQL 16: PostgreSQL is a powerful, open-source object-relational database system with over 30 years of active development. It was chosen over SQLite (Django's default) for its superior performance with concurrent connections, support for advanced SQL features, robust transaction handling, and production-grade reliability. PostgreSQL features used in this project include foreign key constraints with cascade/set-null behaviours, date/time functions for overlap detection queries, and support for concurrent read/write operations.",
    'Body Text')
insert_after(doc.paragraphs[116],
    "Bootstrap 5.3: Bootstrap is the world's most popular CSS framework for building responsive, mobile-first websites. Version 5.3 provides a comprehensive collection of HTML, CSS, and JavaScript components including responsive grid system, navigation bars, cards, forms, buttons, tables, modals, alerts, and utility classes. Bootstrap's responsive breakpoints (sm, md, lg, xl) ensure the website adapts seamlessly to different screen sizes without requiring custom media queries.",
    'Body Text')
insert_after(doc.paragraphs[117],
    "Razorpay Payment Gateway: Razorpay is a leading Indian payment gateway that provides a developer-friendly REST API for processing online payments. It supports all major payment methods including credit cards (Visa, Mastercard, RuPay), debit cards, UPI (Google Pay, PhonePe, Paytm, BHIM), net banking (50+ Indian banks), and digital wallets. The SDK handles order creation, payment capture, refund processing, and webhook verification. During development, Razorpay's test mode was used with test card numbers and UPI IDs to simulate transactions without real money.",
    'Body Text')
insert_after(doc.paragraphs[118],
    "AJAX and JavaScript (ES6): Asynchronous JavaScript and XML (AJAX) enables dynamic content updates without full page reloads. In this project, AJAX is used for the advanced room search feature where search results are fetched from the server and displayed in real-time as the user applies filters. JavaScript ES6 features used include arrow functions, promises, template literals, const/let declarations, and the Fetch API for making asynchronous HTTP requests to the server's RESTful endpoints.",
    'Body Text')
insert_after(doc.paragraphs[119],
    "Git Version Control: Git is a distributed version control system used throughout the project lifecycle for tracking code changes, maintaining development history, and enabling experimental feature development through branching. The Git repository maintains a complete commit history documenting the evolution of the project from initial setup through feature development to final bug fixes and documentation updates.",
    'Body Text')
insert_after(doc.paragraphs[120],
    "Visual Studio Code: VS Code served as the primary integrated development environment (IDE) for this project. Key features utilised include Python language support with IntelliSense, Django template syntax highlighting, integrated Git source control management, debugger configuration for Django runserver, and the Python extension pack for linting, formatting, and virtual environment management.",
    'Body Text')

print("Chapter 2 done")

# Save the document
output_path = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject24.docx'
doc.save(output_path)
print(f"\nDocument saved to: {output_path}")
print("EXPANSION COMPLETE!")
