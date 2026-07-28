"""Replace Data Dictionary plain text with Word tables - uses add_table but reopens doc between tables."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import os

SRC = '/home/nikhil/roombookings3/roombookings/Hoteliers_Myproject/projectreport/nikhilproject241.docx'

doc = Document(SRC)

# Find Data Dictionary boundaries
dd_start = None
dd_end = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == '3.5 Data Dictionary' and 'Heading' in p.style.name:
        dd_start = i
    if dd_start is not None and i > dd_start:
        if '3.6 Entity Relationship' in t and 'Heading' in p.style.name:
            dd_end = i
            break

print(f"DD section: P{dd_start} to P{dd_end-1}")

# Collect data from existing paragraphs
dd_data = {}
current_table = None
intro_para = None
header_para = doc.paragraphs[dd_start]

for i in range(dd_start + 1, dd_end):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if i == dd_start + 1:
        intro_para = p
        continue
    if t.startswith('Table 3.'):
        current_table = t
        dd_data[current_table] = []
    elif current_table and '|' in t:
        parts = [x.strip() for x in t.split('|')]
        dd_data[current_table].append(parts)

print(f"Tables: {list(dd_data.keys())}")
print(f"Total rows: {sum(len(v) for v in dd_data.values())}")

# Collect elements to remove (all between intro and end)
elements_to_remove = []
for i in range(dd_start + 2, dd_end):
    elements_to_remove.append(doc.paragraphs[i]._element)

BODY = doc.element.body
for e in elements_to_remove:
    try:
        BODY.remove(e)
    except:
        pass

print(f"Removed {len(elements_to_remove)} elements")

# Save and reopen for clean state
doc.save(SRC)
doc = Document(SRC)

# Find the intro paragraph again
intro_para = None
for i, p in enumerate(doc.paragraphs):
    if 'defines the structure, data types, constraints' in p.text:
        intro_para = p
        break

if not intro_para:
    print("ERROR: intro paragraph not found!")
    exit(1)

print(f"Intro para at P{i}: {intro_para.text[:60]}")

# Helper functions
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def format_cell(cell, text, bold=False, font_size=9):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    run.bold = bold

def add_table_via_xml(anchor, rows, headers):
    """Create a table using add_table (appended at end), then move it via XML."""
    num_rows = len(rows) + 1
    table = doc.add_table(rows=num_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    col_widths = [Cm(3.5), Cm(3), Cm(4), Cm(6.5)]
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width
    
    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        format_cell(cell, h, bold=True, font_size=9)
        set_cell_shading(cell, '2F5496')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for ri, row_data in enumerate(rows):
        for cj, cell_text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[cj]
            format_cell(cell, cell_text, font_size=9)
            if ri % 2 == 0:
                set_cell_shading(cell, 'D6E4F0')
    
    # Move table after anchor
    table_elem = table._element
    anchor.addnext(table_elem)
    return table_elem

headers = ['Field Name', 'Data Type', 'Constraints', 'Description']

# Save + reopen between each table to keep state clean
def save_and_reopen():
    doc.save(SRC)
    return Document(SRC)

# Track where to insert next
insert_point = intro_para._element

for table_name, rows in dd_data.items():
    # Reopen doc
    doc.save(SRC)
    doc = Document(SRC)
    
    # Find insert_point in new document
    # We stored the table_name text - find it
    # Actually we can't easily persist XML pointers. Let's work differently.
    
    # Find the last paragraph before 3.6 heading
    found_36 = False
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        s = p.style.name
        if t == 'Table 3.10 - Notification':
            # This should be followed by our last table, then 3.6
            pass
    break

# OK this approach is too complicated with reopen.
# Let me do it all in one pass with proper XML insertion.
